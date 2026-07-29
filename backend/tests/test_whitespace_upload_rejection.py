"""Tests for D9: whitespace-only TXT/MD uploads must be rejected, not shipped
as a "valid" 1-page, 0.3s-silence audiobook.

TextExtractor.split_pages always returns at least one page — for
whitespace-only content, _paginate([]) falls through to `[""]`  (see
text_extractor.py's docstring: this is intentional so callers never see an
empty list). That means page_count reads as 1 (not 0) for a whitespace-only
upload, slipping past /audiobook's existing `if page_count == 0` rejection.
The fix extends that check: for non-PDF uploads with page_count == 1, the
lone page's actual (stripped) text must be non-empty too.
"""

from __future__ import annotations

import io
import shutil
import tempfile

import pytest

from app.services.audiobook_store import AudiobookStore


@pytest.fixture(autouse=True)
def isolated_audiobooks_dir(monkeypatch):
    tmp = tempfile.mkdtemp(prefix="ss_audiobooks_test_")

    class _PatchedSettings:
        @property
        def AUDIOBOOKS_DIR(self) -> str:
            return tmp

    monkeypatch.setattr("app.services.audiobook_store.settings", _PatchedSettings())
    AudiobookStore._reset_for_tests()
    yield tmp
    AudiobookStore._reset_for_tests()
    shutil.rmtree(tmp, ignore_errors=True)


def _upload(client, filename: str, content: bytes):
    return client.post(
        "/audiobook",
        files={"file": (filename, io.BytesIO(content), "text/plain")},
    )


def test_whitespace_only_txt_upload_is_rejected():
    from fastapi.testclient import TestClient

    from app.main import app

    client = TestClient(app)
    resp = _upload(client, "blank.txt", b"   \n\n\t  \n   ")

    assert resp.status_code == 400
    assert "no extractable content" in resp.json()["detail"]
    # No book row must survive a rejected upload.
    assert AudiobookStore.list_books() == []


def test_whitespace_only_md_upload_is_rejected():
    from fastapi.testclient import TestClient

    from app.main import app

    client = TestClient(app)
    resp = _upload(client, "blank.md", b"\n\n   \n\n  \n")

    assert resp.status_code == 400
    assert "no extractable content" in resp.json()["detail"]
    assert AudiobookStore.list_books() == []


def test_genuinely_empty_txt_upload_is_still_rejected():
    """Non-regression: the pre-existing zero-byte-content rejection (a
    different, earlier check than the D9 fix) must still work."""
    from fastapi.testclient import TestClient

    from app.main import app

    client = TestClient(app)
    resp = _upload(client, "empty.txt", b"")

    assert resp.status_code == 400
    assert "empty" in resp.json()["detail"].lower()


def test_txt_upload_with_real_content_still_succeeds():
    """Non-regression: a real single-page TXT upload (page_count == 1, but
    genuinely non-blank) must not be caught by the D9 guard."""
    from fastapi.testclient import TestClient

    from app.main import app

    client = TestClient(app)
    resp = _upload(client, "note.txt", b"This is a real, short note.")

    assert resp.status_code == 200
    body = resp.json()
    assert body["page_count"] == 1
    assert AudiobookStore.read_meta(body["book_id"]) is not None


def test_multi_page_txt_upload_unaffected_by_d9_guard():
    """Non-regression: the D9 guard only triggers for page_count == 1 — a
    multi-page upload must never be routed through it."""
    from fastapi.testclient import TestClient

    from app.main import app

    client = TestClient(app)
    # Two real, distinct paragraphs separated by a blank line.
    content = ("Page one content. " * 50 + "\n\n" + "Page two content. " * 400).encode()
    resp = _upload(client, "book.txt", content)

    assert resp.status_code == 200
    assert resp.json()["page_count"] >= 2


def test_whitespace_only_docx_upload_is_rejected():
    """The D9 guard's `not is_pdf` condition covers DOCX too, not just
    TXT/MD — TextExtractor.read_text() strips DOCX down to only non-blank
    paragraphs, so a DOCX containing exclusively empty/whitespace paragraphs
    hits the exact same single-blank-page synthetic-pagination edge case."""
    from docx import Document
    from fastapi.testclient import TestClient

    from app.main import app

    doc = Document()
    doc.add_paragraph("   ")
    doc.add_paragraph("")
    buf = io.BytesIO()
    doc.save(buf)

    client = TestClient(app)
    resp = client.post(
        "/audiobook",
        files={
            "file": (
                "blank.docx",
                io.BytesIO(buf.getvalue()),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert resp.status_code == 400
    assert "no extractable content" in resp.json()["detail"]
    assert AudiobookStore.list_books() == []
