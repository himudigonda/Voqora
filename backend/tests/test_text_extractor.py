"""Tests for TextExtractor — TXT / DOCX / Markdown → audiobook page pipeline.

No mocking needed: TextExtractor reads real text files written to temp dirs.
Coverage:
  - split_pages(): empty, short (one page), multi-page at paragraph boundaries,
    CRLF normalization, short paragraph not split mid-page
  - read_text(): UTF-8 TXT, Latin-1 replacement
  - is_image_only(): always False
  - read_outline(): None for TXT (no native structure); real DOCX Heading-N
    styles and Markdown #/## headers for those formats (T3.2/T3.3), each
    position cross-checked against split_pages()'s actual pagination so a
    heading's reported start_page always matches where the pipeline puts it
  - page_count(): delegates to read_text + split_pages
  - sample_word_count() / sample_char_count(): three-page sampling
  - extract_one(): idempotent; writes ALL pages on first call
  - render_cover(): creates a JPEG; idempotent
  - _atomic_write() / _atomic_write_bytes(): tmp+rename atomicity
"""

from __future__ import annotations

import os
import shutil
import time

import pytest

from app.services.text_extractor import _WORDS_PER_PAGE, TextExtractor

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture()
def tmp_book(tmp_path, monkeypatch):
    """Redirect AudiobookStore to tmp_path; yield (book_id, source_path, tmp_path)."""
    from app.services.audiobook_store import AudiobookStore

    class _S:
        @property
        def AUDIOBOOKS_DIR(self):
            return str(tmp_path)

    monkeypatch.setattr("app.services.audiobook_store.settings", _S())
    AudiobookStore._reset_for_tests()

    bid = AudiobookStore.create_book("sample.txt")
    meta = AudiobookStore.initial_meta(
        bid, "sample.txt", 1, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
    )
    meta["file_ext"] = "txt"
    AudiobookStore.write_meta(bid, meta)

    source = AudiobookStore.source_file_path(bid, "txt")
    os.makedirs(os.path.dirname(source), exist_ok=True)

    yield bid, source, tmp_path

    AudiobookStore._reset_for_tests()
    shutil.rmtree(str(tmp_path), ignore_errors=True)


# ---------------------------------------------------------------------------
# split_pages
# ---------------------------------------------------------------------------


def test_split_pages_empty_text_returns_single_empty_string():
    assert TextExtractor.split_pages("") == [""]


def test_split_pages_whitespace_only_returns_single_empty_string():
    # No meaningful paragraphs → [""]
    assert TextExtractor.split_pages("   \n\n   ") == [""]


def test_split_pages_short_text_is_one_page():
    text = "Short paragraph.\n\nAnother short paragraph."
    pages = TextExtractor.split_pages(text)
    assert len(pages) == 1
    assert "Short paragraph" in pages[0]
    assert "Another short paragraph" in pages[0]


def test_split_pages_splits_at_paragraph_boundary():
    """A paragraph that pushes total over _WORDS_PER_PAGE triggers a new page."""
    big = " ".join(f"word{i}" for i in range(_WORDS_PER_PAGE + 1))
    small = "This is page 2."
    text = big + "\n\n" + small
    pages = TextExtractor.split_pages(text)
    assert len(pages) == 2
    assert "page 2" in pages[1]


def test_split_pages_single_oversized_paragraph_falls_back_to_word_chunking():
    """No double-newline → no paragraph split point → falls back to chunking
    by raw word count so a long unbroken block still yields multiple pages."""
    single_para = " ".join(f"w{i}" for i in range(_WORDS_PER_PAGE + 50))
    pages = TextExtractor.split_pages(single_para)
    assert len(pages) == 2
    assert len(pages[0].split()) == _WORDS_PER_PAGE
    assert len(pages[1].split()) == 50


def test_split_pages_word_chunk_fallback_preserves_all_words():
    """The no-blank-line fallback must not drop or duplicate any word."""
    words = [f"w{i}" for i in range(_WORDS_PER_PAGE * 3 + 7)]
    text = " ".join(words)
    pages = TextExtractor.split_pages(text)
    assert len(pages) == 4
    rejoined = " ".join(pages).split()
    assert rejoined == words


def test_split_pages_three_equal_pages():
    para = " ".join(f"word{i}" for i in range(_WORDS_PER_PAGE))
    text = f"{para}\n\n{para}\n\n{para}"
    pages = TextExtractor.split_pages(text)
    assert len(pages) == 3


def test_split_pages_crlf_normalised():
    text = "Para one.\r\n\r\nPara two."
    pages = TextExtractor.split_pages(text)
    # Both paras fit on one page
    assert len(pages) == 1
    assert "Para one" in pages[0]
    assert "Para two" in pages[0]


def test_split_pages_cr_only_normalised():
    text = "Para A.\r\rPara B."
    pages = TextExtractor.split_pages(text)
    assert len(pages) == 1
    assert "Para A" in pages[0]


def test_split_pages_preserves_all_words():
    """No word is lost during splitting."""
    paras = [" ".join(f"x{j}" for j in range(_WORDS_PER_PAGE // 2)) for _ in range(6)]
    text = "\n\n".join(paras)
    pages = TextExtractor.split_pages(text)
    joined = " ".join(pages)
    for para in paras:
        for word in para.split():
            assert word in joined


# ---------------------------------------------------------------------------
# read_text
# ---------------------------------------------------------------------------


def test_read_text_utf8_txt(tmp_path):
    p = tmp_path / "doc.txt"
    p.write_text("Hello, world!", encoding="utf-8")
    assert TextExtractor.read_text(str(p)) == "Hello, world!"


def test_read_text_multiline(tmp_path):
    content = "Line 1\nLine 2\nLine 3"
    p = tmp_path / "doc.txt"
    p.write_text(content, encoding="utf-8")
    assert TextExtractor.read_text(str(p)) == content


def test_read_text_latin1_falls_back_to_replacement(tmp_path):
    p = tmp_path / "doc.txt"
    p.write_bytes(b"caf\xe9")  # 'café' in latin-1
    result = TextExtractor.read_text(str(p))
    assert "caf" in result  # partial content preserved, no crash


def test_read_text_md_treated_as_plain(tmp_path):
    p = tmp_path / "README.md"
    p.write_text("# Title\n\nBody text.", encoding="utf-8")
    assert "Title" in TextExtractor.read_text(str(p))


# ---------------------------------------------------------------------------
# is_image_only / read_outline
# ---------------------------------------------------------------------------


def test_is_image_only_always_false():
    assert TextExtractor.is_image_only("/any/path") is False


def test_read_outline_always_none():
    assert TextExtractor.read_outline("/any/path") is None


def test_read_outline_returns_none_for_txt_extension(tmp_path):
    """Plain TXT has no native structural signal — always falls through to
    Gemini/fallback section detection, same as before this sprint."""
    p = tmp_path / "plain.txt"
    p.write_text("Just some text.", encoding="utf-8")
    assert TextExtractor.read_outline(str(p)) is None


# ---------------------------------------------------------------------------
# read_outline — DOCX heading-style detection (T3.2)
# ---------------------------------------------------------------------------


def test_read_outline_docx_detects_heading_styles(tmp_path):
    """Heading-1/Heading-2-styled paragraphs become structured heading
    entries (title, start_page, level) — plain "Normal"-styled paragraphs
    are not structural boundaries."""
    from docx import Document

    doc = Document()
    doc.add_heading("Chapter One", level=1)
    doc.add_paragraph("Body text for chapter one.")
    doc.add_heading("Section 1.1", level=2)
    doc.add_paragraph("More body text here.")
    doc.add_heading("Chapter Two", level=1)
    doc.add_paragraph("Final body text.")
    path = str(tmp_path / "headings.docx")
    doc.save(path)

    headings = TextExtractor.read_outline(path)

    assert headings == [
        {"title": "Chapter One", "start_page": 1, "level": 1},
        {"title": "Section 1.1", "start_page": 1, "level": 2},
        {"title": "Chapter Two", "start_page": 1, "level": 1},
    ]


def test_read_outline_docx_with_no_heading_styles_returns_none(tmp_path):
    """A DOCX with zero heading-styled paragraphs falls through cleanly to
    Gemini/fallback, exactly as it does today."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Just a plain paragraph.")
    doc.add_paragraph("Another plain paragraph, still no heading styles.")
    path = str(tmp_path / "no_headings.docx")
    doc.save(path)

    assert TextExtractor.read_outline(path) is None


def test_read_outline_docx_heading_position_matches_real_pagination(tmp_path):
    """A heading after a page-overflowing paragraph must report the SAME
    start_page split_pages() actually places it on — the whole point of
    routing heading detection through the shared _paginate() helper instead
    of a separately-derived position that could silently drift out of sync."""
    from docx import Document

    doc = Document()
    big_para = " ".join(f"word{i}" for i in range(_WORDS_PER_PAGE + 20))
    doc.add_paragraph(big_para)
    doc.add_heading("Chapter Two Starts Here", level=1)
    doc.add_paragraph("Body text.")
    path = str(tmp_path / "heading_page2.docx")
    doc.save(path)

    headings = TextExtractor.read_outline(path)
    assert headings == [
        {"title": "Chapter Two Starts Here", "start_page": 2, "level": 1}
    ]

    # Cross-check against the real pagination the pipeline will actually use.
    pages = TextExtractor.split_pages(TextExtractor.read_text(path))
    assert len(pages) == 2
    assert "Chapter Two Starts Here" in pages[1]
    assert "Chapter Two Starts Here" not in pages[0]


def test_read_outline_docx_heading_position_survives_embedded_double_break(tmp_path):
    """Regression: a SINGLE docx paragraph whose own .text contains an
    embedded blank-line run (e.g. two consecutive manual line breaks —
    python-docx renders each <w:br/> as "\\n", so two in a row is "\\n\\n")
    must not desync heading position from real pagination.

    The real pipeline paginates read_text()'s "\\n\\n".join(...) of every
    paragraph, RE-SPLIT on \\n{2,} by _split_into_paragraphs() — so this one
    docx paragraph becomes TWO synthetic paragraphs in the real pipeline.
    A heading-position implementation that instead paginates over raw
    doc.paragraphs objects directly (one entry per object, ignoring an
    internal \\n{2,} run) sees a different, shorter list and reports the
    heading one page too EARLY — silently pointing at a page that doesn't
    even contain the heading text, and downstream corrupts the section's
    page_to_time lookup (falls back to 0.0 — seeks to the very start of the
    book) even though the title itself is correct.

    Verified by hand before writing this test: with 50 words, an embedded
    double line-break, then 400 more words, then this heading, the OLD
    (pre-fix) pagination-over-raw-paragraphs approach reports start_page=2;
    the real pipeline (and the fix) report start_page=3 — and only page 3
    actually contains the heading text (see the cross-check below).
    """
    from docx import Document

    doc = Document()
    p = doc.add_paragraph()
    p.add_run(" ".join(f"a{i}" for i in range(50)))
    br_run = p.add_run()
    br_run.add_break()
    br_run.add_break()
    p.add_run(" ".join(f"b{i}" for i in range(400)))
    assert "\n\n" in p.text, "fixture must actually contain an embedded \\n\\n"

    doc.add_heading("Real Chapter Two", level=1)
    doc.add_paragraph("Body text after the heading.")
    path = str(tmp_path / "embedded_double_break.docx")
    doc.save(path)

    headings = TextExtractor.read_outline(path)
    assert headings == [{"title": "Real Chapter Two", "start_page": 3, "level": 1}]

    # Ground truth: the heading's reported page must be where the pipeline
    # ACTUALLY puts it — not merely "a page", the RIGHT page.
    pages = TextExtractor.split_pages(TextExtractor.read_text(path))
    assert len(pages) == 3
    assert "Real Chapter Two" in pages[2]
    assert "Real Chapter Two" not in pages[0]
    assert "Real Chapter Two" not in pages[1]


def test_read_outline_docx_heading_style_without_number_defaults_to_level_1(tmp_path):
    """A style literally named "Heading" (no numeric suffix) still counts as
    structural per spec ("style.name starting with Heading"), defaulting to
    level 1 instead of crashing on int('')."""
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document()
    custom = doc.styles.add_style("Heading", WD_STYLE_TYPE.PARAGRAPH)
    p = doc.add_paragraph("Untitled-level heading")
    p.style = custom
    path = str(tmp_path / "custom_heading.docx")
    doc.save(path)

    headings = TextExtractor.read_outline(path)
    assert headings == [
        {"title": "Untitled-level heading", "start_page": 1, "level": 1}
    ]


def test_read_outline_docx_returns_none_on_corrupt_or_missing_file(tmp_path):
    """Mirrors PDFExtractor.read_outline's resilience: a corrupt/missing DOCX
    must degrade to None (falling through to Gemini/fallback), not crash the
    whole book's processing."""
    missing = str(tmp_path / "does_not_exist.docx")
    assert TextExtractor.read_outline(missing) is None

    corrupt = str(tmp_path / "corrupt.docx")
    with open(corrupt, "wb") as f:
        f.write(b"not a real docx, just some bytes")
    assert TextExtractor.read_outline(corrupt) is None


# ---------------------------------------------------------------------------
# read_outline — Markdown #/## header detection (T3.3)
# ---------------------------------------------------------------------------


def test_read_outline_md_detects_h1_and_h2_headers(tmp_path):
    content = (
        "# Chapter One\n\n"
        "Body text for chapter one.\n\n"
        "## Section 1.1\n\n"
        "More body text here.\n\n"
        "# Chapter Two\n\n"
        "Final body text.\n"
    )
    path = tmp_path / "headers.md"
    path.write_text(content, encoding="utf-8")

    headings = TextExtractor.read_outline(str(path))

    assert headings == [
        {"title": "Chapter One", "start_page": 1, "level": 1},
        {"title": "Section 1.1", "start_page": 1, "level": 2},
        {"title": "Chapter Two", "start_page": 1, "level": 1},
    ]


def test_read_outline_md_with_no_headers_returns_none(tmp_path):
    """Plain MD with no #/## headers falls through cleanly, same contract as
    a DOCX with no heading styles."""
    content = "Just a plain paragraph.\n\nAnother plain paragraph, no headers.\n"
    path = tmp_path / "plain.md"
    path.write_text(content, encoding="utf-8")

    assert TextExtractor.read_outline(str(path)) is None


def test_read_outline_md_ignores_h3_and_deeper(tmp_path):
    """Only #/## count as section boundaries per spec — a document with only
    deeper headers has no *structural* signal and must return None (deeper
    headers are sub-structure within a section, not a chapter break)."""
    content = (
        "Just a plain paragraph with no top-level headers.\n\n"
        "### This is a level-3 header and must be ignored.\n\n"
        "Another paragraph.\n"
    )
    path = tmp_path / "h3_only.md"
    path.write_text(content, encoding="utf-8")

    assert TextExtractor.read_outline(str(path)) is None


def test_read_outline_md_ignores_hash_without_space():
    """ATX syntax requires a space after the leading hashes — '#nospace' or
    an inline '#hashtag' must not be misdetected as a header."""
    assert TextExtractor._parse_md_header("#nospace") is None
    assert TextExtractor._parse_md_header("some #hashtag inline text") is None
    assert TextExtractor._parse_md_header("####### too many hashes") is None
    assert TextExtractor._parse_md_header("# Valid Header") == (1, "Valid Header")
    assert TextExtractor._parse_md_header("## Valid H2 ##") == (2, "Valid H2")


def test_read_outline_md_preserves_trailing_hash_that_is_not_a_closer():
    """Regression: a title that itself legitimately ends in '#' (e.g. a
    programming-language name) must not have that character silently eaten
    by the optional-ATX-closing-sequence stripping. Per CommonMark, a
    trailing '#' run only counts as a closer when preceded by whitespace —
    "Learning F#" has no space before its trailing '#', so it must survive
    intact, while "## Title ##" (space before the closing hashes) still has
    its closer correctly stripped."""
    assert TextExtractor._parse_md_header("# Learning F#") == (1, "Learning F#")
    assert TextExtractor._parse_md_header("## C# and F#") == (2, "C# and F#")
    # The legitimate closer case must still work after the fix.
    assert TextExtractor._parse_md_header("## Title ##") == (2, "Title")
    assert TextExtractor._parse_md_header("# Single Closer #") == (1, "Single Closer")


def test_read_outline_md_header_position_matches_real_pagination(tmp_path):
    """Same cross-check as the DOCX case: a header after a page-overflowing
    paragraph must land on the page split_pages() actually puts it on."""
    big_para = " ".join(f"word{i}" for i in range(_WORDS_PER_PAGE + 20))
    content = f"{big_para}\n\n# Chapter Two Starts Here\n\nBody text.\n"
    path = tmp_path / "heading_page2.md"
    path.write_text(content, encoding="utf-8")

    headings = TextExtractor.read_outline(str(path))
    assert headings == [
        {"title": "Chapter Two Starts Here", "start_page": 2, "level": 1}
    ]

    pages = TextExtractor.split_pages(TextExtractor.read_text(str(path)))
    assert len(pages) == 2
    assert "Chapter Two Starts Here" in pages[1]
    assert "Chapter Two Starts Here" not in pages[0]


def test_read_outline_md_returns_none_on_missing_file(tmp_path):
    """Mirrors PDFExtractor.read_outline's resilience: a missing/unreadable
    MD file must degrade to None, not crash the whole book's processing."""
    missing = str(tmp_path / "does_not_exist.md")
    assert TextExtractor.read_outline(missing) is None


# ---------------------------------------------------------------------------
# page_count
# ---------------------------------------------------------------------------


def test_page_count_single_page(tmp_path):
    p = tmp_path / "short.txt"
    p.write_text("Only a few words here.", encoding="utf-8")
    assert TextExtractor.page_count(str(p)) == 1


def test_page_count_multi_page(tmp_path):
    para = " ".join(f"w{i}" for i in range(_WORDS_PER_PAGE + 1))
    text = para + "\n\n" + para
    p = tmp_path / "long.txt"
    p.write_text(text, encoding="utf-8")
    assert TextExtractor.page_count(str(p)) >= 2


def test_page_count_empty_file(tmp_path):
    p = tmp_path / "empty.txt"
    p.write_text("", encoding="utf-8")
    assert TextExtractor.page_count(str(p)) == 1


# ---------------------------------------------------------------------------
# sample_word_count / sample_char_count
# ---------------------------------------------------------------------------


def test_sample_word_count_single_page(tmp_path):
    p = tmp_path / "doc.txt"
    p.write_text("one two three", encoding="utf-8")
    assert TextExtractor.sample_word_count(str(p)) == 3


def test_sample_char_count_single_page(tmp_path):
    content = "hello world"
    p = tmp_path / "doc.txt"
    p.write_text(content, encoding="utf-8")
    assert TextExtractor.sample_char_count(str(p)) == len(content)


def test_sample_word_count_empty_returns_zero(tmp_path):
    p = tmp_path / "empty.txt"
    p.write_text("", encoding="utf-8")
    assert TextExtractor.sample_word_count(str(p)) == 0


def test_sample_char_count_empty_returns_zero(tmp_path):
    p = tmp_path / "empty.txt"
    p.write_text("", encoding="utf-8")
    assert TextExtractor.sample_char_count(str(p)) == 0


# ---------------------------------------------------------------------------
# extract_one
# ---------------------------------------------------------------------------


def test_extract_one_writes_page_file(tmp_book):
    bid, source, _ = tmp_book
    from app.services.audiobook_store import AudiobookStore

    with open(source, "w", encoding="utf-8") as f:
        f.write("Hello world.\n\nSecond paragraph.")

    TextExtractor.extract_one(bid, 1)

    out = AudiobookStore.page_raw_path(bid, 1)
    assert os.path.exists(out)
    assert len(open(out, encoding="utf-8").read()) > 0


def test_extract_one_is_idempotent(tmp_book):
    """Second call for an existing file must not overwrite it."""
    bid, source, _ = tmp_book
    from app.services.audiobook_store import AudiobookStore

    with open(source, "w", encoding="utf-8") as f:
        f.write("One paragraph.\n\nTwo paragraph content here.")

    TextExtractor.extract_one(bid, 1)
    out = AudiobookStore.page_raw_path(bid, 1)
    mtime1 = os.path.getmtime(out)

    time.sleep(0.02)
    TextExtractor.extract_one(bid, 1)
    mtime2 = os.path.getmtime(out)

    assert mtime1 == mtime2, "extract_one must not overwrite existing page file"


def test_extract_one_writes_all_pages_on_first_call(tmp_book):
    """First extract_one() call writes ALL pages so later calls skip I/O."""
    bid, source, _ = tmp_book
    from app.services.audiobook_store import AudiobookStore

    # Update meta to reflect 2 pages
    meta = AudiobookStore.read_meta(bid) or {}
    meta["page_count"] = 2
    AudiobookStore.write_meta(bid, meta)

    para = " ".join(f"word{i}" for i in range(_WORDS_PER_PAGE + 1))
    with open(source, "w", encoding="utf-8") as f:
        f.write(para + "\n\nSecond page content here.")

    # Only request page 1
    TextExtractor.extract_one(bid, 1)

    # Page 2 must also exist (written proactively)
    out2 = AudiobookStore.page_raw_path(bid, 2)
    assert os.path.exists(out2), "extract_one(bid, 1) must pre-write page 2"


def test_extract_one_content_non_empty(tmp_book):
    bid, source, _ = tmp_book
    from app.services.audiobook_store import AudiobookStore

    with open(source, "w", encoding="utf-8") as f:
        f.write("Some content on page one.")

    TextExtractor.extract_one(bid, 1)
    content = open(AudiobookStore.page_raw_path(bid, 1), encoding="utf-8").read()
    assert len(content.strip()) > 0


# ---------------------------------------------------------------------------
# render_cover
# ---------------------------------------------------------------------------


def test_render_cover_creates_jpeg(tmp_book):
    bid, _, _ = tmp_book
    from app.services.audiobook_store import AudiobookStore

    TextExtractor.render_cover(bid)

    cover = AudiobookStore.cover_path(bid)
    assert os.path.exists(cover), "cover.jpg must be created by render_cover"
    with open(cover, "rb") as f:
        magic = f.read(2)
    assert magic == b"\xff\xd8", "cover must be a valid JPEG (FF D8 magic bytes)"


def test_render_cover_is_idempotent(tmp_book):
    """Second call must not overwrite an existing cover."""
    bid, _, _ = tmp_book
    from app.services.audiobook_store import AudiobookStore

    TextExtractor.render_cover(bid)
    cover = AudiobookStore.cover_path(bid)
    mtime1 = os.path.getmtime(cover)

    time.sleep(0.02)
    TextExtractor.render_cover(bid)
    mtime2 = os.path.getmtime(cover)

    assert mtime1 == mtime2, "render_cover must not overwrite existing cover"


def test_render_cover_is_readable_jpeg(tmp_book):
    """The generated cover must be decodeable by Pillow (not corrupted)."""
    bid, _, _ = tmp_book
    from PIL import Image

    from app.services.audiobook_store import AudiobookStore

    TextExtractor.render_cover(bid)
    cover = AudiobookStore.cover_path(bid)
    img = Image.open(cover)
    assert img.mode == "RGB"
    assert img.width > 0 and img.height > 0


# ---------------------------------------------------------------------------
# _atomic_write / _atomic_write_bytes
# ---------------------------------------------------------------------------


def test_atomic_write_creates_nested_dirs_and_file(tmp_path):
    path = str(tmp_path / "sub" / "dir" / "out.txt")
    TextExtractor._atomic_write(path, "hello")
    assert open(path).read() == "hello"
    assert not os.path.exists(path + ".tmp"), "tmp file must be cleaned up"


def test_atomic_write_replaces_existing(tmp_path):
    path = str(tmp_path / "out.txt")
    with open(path, "w") as f:
        f.write("old")
    TextExtractor._atomic_write(path, "new")
    assert open(path).read() == "new"


def test_atomic_write_bytes_creates_file(tmp_path):
    path = str(tmp_path / "out.bin")
    TextExtractor._atomic_write_bytes(path, b"\x01\x02\x03")
    assert open(path, "rb").read() == b"\x01\x02\x03"
    assert not os.path.exists(path + ".tmp")


def test_atomic_write_bytes_replaces_existing(tmp_path):
    path = str(tmp_path / "out.bin")
    with open(path, "wb") as f:
        f.write(b"\xff\xff")
    TextExtractor._atomic_write_bytes(path, b"\x00")
    assert open(path, "rb").read() == b"\x00"


def test_split_pages_oversized_paragraph_amid_normal_paragraphs_still_chunks():
    """A single huge paragraph followed by a short one must not collapse into
    one giant page just because it isn't the *only* paragraph in the doc."""
    huge = " ".join(f"w{i}" for i in range(_WORDS_PER_PAGE * 5))
    text = huge + "\n\nThe End."
    pages = TextExtractor.split_pages(text)
    assert len(pages) > 1
    assert all(len(p.split()) <= _WORDS_PER_PAGE for p in pages[:-1])
    assert "The End" in pages[-1]
    # No word lost or duplicated.
    rejoined = " ".join(pages).split()
    assert rejoined == huge.split() + ["The", "End."]
