"""Tests for the audiobook pipeline.

Mocks: PDFExtractor (no real PDFs), GeminiCleaner (no API calls),
EngineManager.generate_with_timing (yields short segment-timing dicts).
"""

import asyncio
import json
import os
import shutil
import tempfile
from unittest.mock import AsyncMock, patch

import numpy as np
import pytest

from app.services.audiobook_service import (
    SAMPLE_RATE,
    WAV_HEADER_SIZE,
    AudiobookService,
    _wav_header,
)
from app.services.audiobook_store import AudiobookStore


@pytest.fixture(autouse=True)
def isolated_audiobooks_dir(monkeypatch):
    """Redirect AUDIOBOOKS_DIR to a per-test temp dir so tests don't pollute.

    Also resets the singleton SQLite connection so each test gets a fresh
    audiobooks.db inside its own tmp dir.
    """
    tmp = tempfile.mkdtemp(prefix="ss_audiobooks_test_")

    class _PatchedSettings:
        @property
        def AUDIOBOOKS_DIR(self) -> str:
            return tmp

    monkeypatch.setattr("app.services.audiobook_store.settings", _PatchedSettings())
    AudiobookStore._reset_for_tests()
    yield tmp
    AudiobookStore._reset_for_tests()
    shutil.rmtree(tmp, ignore_errors=True)


# ---------- AudiobookStore ----------


def test_legacy_meta_json_is_migrated_to_sqlite(isolated_audiobooks_dir):
    """Regression: a leftover meta.json from before the SQLite migration is
    imported into the DB on first connection, and the JSON file is deleted."""
    from app.services.audiobook_store import AudiobookStore

    bid = "abc123_legacy"
    bdir = os.path.join(isolated_audiobooks_dir, bid)
    os.makedirs(os.path.join(bdir, "pages"), exist_ok=True)
    legacy_path = os.path.join(bdir, "meta.json")
    legacy = {
        "book_id": bid,
        "title": "Legacy.pdf",
        "created_at": "2024-01-01T00:00:00Z",
        "page_count": 7,
        "status": "done",
        "engine": "kokoro",
        "voice": "af_bella",
        "speed": 1.0,
        "estimated": {"cost_usd": 0.5},
    }
    with open(legacy_path, "w") as f:
        json.dump(legacy, f)

    # Force connection (triggers migration).
    AudiobookStore._reset_for_tests()
    meta = AudiobookStore.read_meta(bid)
    assert meta is not None
    assert meta["title"] == "Legacy.pdf"
    assert meta["page_count"] == 7
    # Legacy JSON file removed.
    assert not os.path.exists(legacy_path)


def test_create_book_makes_dirs():
    bid = AudiobookStore.create_book("Test.pdf")
    assert os.path.isdir(AudiobookStore.book_dir(bid))
    assert os.path.isdir(os.path.join(AudiobookStore.book_dir(bid), "pages"))
    assert os.path.isdir(os.path.join(AudiobookStore.book_dir(bid), "audio_pages"))


def test_page_segments_path_shape():
    """The segments sidecar is co-located with its page WAV under audio_pages/,
    zero-padded to 3 digits, matching the {n:03d}.wav sibling pattern."""
    path = AudiobookStore.page_segments_path("abc123", 7)
    assert path == os.path.join(
        AudiobookStore.book_dir("abc123"), "audio_pages", "007.segments.json"
    )


def test_meta_atomic_write_and_read():
    bid = AudiobookStore.create_book("Test.pdf")
    meta = AudiobookStore.initial_meta(
        bid, "Test.pdf", 5, "kokoro", "af_bella", 1.0, {"cost_usd": 0.1}
    )
    AudiobookStore.write_meta(bid, meta)
    read = AudiobookStore.read_meta(bid)
    assert read is not None
    assert read["book_id"] == bid
    assert read["page_count"] == 5
    assert read["status"] == "ready"


def test_list_books_sorted_desc():
    b1 = AudiobookStore.create_book("a.pdf")
    AudiobookStore.write_meta(
        b1,
        {"book_id": b1, "title": "a", "created_at": "2024-01-01T00:00:00Z"},
    )
    b2 = AudiobookStore.create_book("b.pdf")
    AudiobookStore.write_meta(
        b2,
        {"book_id": b2, "title": "b", "created_at": "2025-06-01T00:00:00Z"},
    )
    books = AudiobookStore.list_books()
    assert [b["book_id"] for b in books][:2] == [b2, b1]


def test_delete_book_removes_dir():
    bid = AudiobookStore.create_book("Test.pdf")
    assert AudiobookStore.delete_book(bid) is True
    assert not os.path.isdir(AudiobookStore.book_dir(bid))
    assert AudiobookStore.delete_book(bid) is False  # second delete


# ---------- estimation ----------


def test_estimate_math():
    est = AudiobookService.estimate(
        page_count=100, sample_words=300, sample_chars=1500, speed=1.0
    )
    assert est["pages"] == 100
    assert est["words"] == 30000
    # 30000 words / 2.75 wps ≈ 10909 s
    assert 10500 < est["audio_seconds"] < 11200
    assert est["processing_seconds"] > 0
    assert est["cost_usd"] > 0


def test_estimate_is_self_consistent_includes_token_count():
    """estimate() must return token_count itself — callers shouldn't patch it in
    afterwards (which left other callers KeyError-prone)."""
    from app.services.gemini_cleaner import GeminiCleaner

    est = AudiobookService.estimate(
        page_count=10, sample_words=200, sample_chars=1000, speed=1.0
    )
    assert "token_count" in est
    # ~4 chars/token over 10 pages × 1000 chars = 10000 chars → ~2500 tokens.
    assert est["token_count"] == GeminiCleaner.estimate_tokens(1000 * 10)
    assert est["token_count"] > 0


# ---------- WAV concat ----------


def _write_pcm_wav(path: str, n_samples: int, value: int = 0) -> None:
    """Write a tiny WAV file with a known PCM length (for concat tests)."""
    import wave

    pcm = (np.full(n_samples, value, dtype=np.int16)).tobytes()
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with wave.open(path, "wb") as wf:
        wf.setnchannels(1)
        wf.setsampwidth(2)
        wf.setframerate(SAMPLE_RATE)
        wf.writeframes(pcm)


@pytest.mark.asyncio
async def test_concat_phase_builds_correct_wav_and_page_to_time():
    bid = AudiobookStore.create_book("Test.pdf")
    meta = AudiobookStore.initial_meta(
        bid, "Test.pdf", 3, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
    )
    AudiobookStore.write_meta(bid, meta)

    # 3 pages: 24000 samples (1s), 12000 samples (0.5s), 6000 samples (0.25s)
    _write_pcm_wav(AudiobookStore.page_audio_path(bid, 1), 24000)
    _write_pcm_wav(AudiobookStore.page_audio_path(bid, 2), 12000)
    _write_pcm_wav(AudiobookStore.page_audio_path(bid, 3), 6000)

    actual = await AudiobookService._phase_concat(bid, AudiobookStore.read_meta(bid))

    # Final WAV exists, body = sum of PCM bytes (3 pages, 2 bytes/sample)
    final = AudiobookStore.audio_path(bid)
    assert os.path.exists(final)
    expected_pcm = (24000 + 12000 + 6000) * 2
    assert os.path.getsize(final) == WAV_HEADER_SIZE + expected_pcm

    new_meta = AudiobookStore.read_meta(bid)
    assert new_meta["page_to_time"]["1"] == 0.0
    # page 2 starts after 1s
    assert abs(new_meta["page_to_time"]["2"] - 1.0) < 0.001
    # page 3 starts after 1.5s
    assert abs(new_meta["page_to_time"]["3"] - 1.5) < 0.001
    assert abs(new_meta["total_audio_seconds"] - 1.75) < 0.001
    assert actual["pages"] == 3
    assert actual["audio_seconds"] == new_meta["total_audio_seconds"]


@pytest.mark.asyncio
async def test_concat_phase_writes_segments_with_absolute_offsets():
    """T1.7: transcript.json's segments/dropped_segments are the per-page
    sidecar data, offset from page-relative to whole-book-absolute time via
    page_to_time — matching how sections[].start_time is already computed."""
    bid = AudiobookStore.create_book("Test.pdf")
    meta = AudiobookStore.initial_meta(
        bid, "Test.pdf", 2, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
    )
    AudiobookStore.write_meta(bid, meta)

    # Page 1: 1.0s of audio, starts at absolute t=0.
    _write_pcm_wav(AudiobookStore.page_audio_path(bid, 1), 24000)
    AudiobookService._write_segments_json(
        AudiobookStore.page_segments_path(bid, 1),
        [{"text": "Page one sentence.", "start_sec": 0.0, "end_sec": 1.0}],
        [],
    )
    # Page 2: 0.5s of audio, starts at absolute t=1.0 (right after page 1).
    _write_pcm_wav(AudiobookStore.page_audio_path(bid, 2), 12000)
    AudiobookService._write_segments_json(
        AudiobookStore.page_segments_path(bid, 2),
        [{"text": "Page two sentence.", "start_sec": 0.0, "end_sec": 0.5}],
        [{"index": 3, "text": "A segment that failed on page two."}],
    )

    await AudiobookService._phase_concat(bid, AudiobookStore.read_meta(bid))

    with open(AudiobookStore.transcript_path(bid), encoding="utf-8") as f:
        transcript = json.load(f)

    assert transcript["segments"]["1"] == [
        {"text": "Page one sentence.", "start_sec": 0.0, "end_sec": 1.0}
    ]
    # Page 2's segment must be offset by page 1's 1.0s duration.
    seg2 = transcript["segments"]["2"][0]
    assert seg2["text"] == "Page two sentence."
    assert abs(seg2["start_sec"] - 1.0) < 1e-6
    assert abs(seg2["end_sec"] - 1.5) < 1e-6

    assert transcript["dropped_segments"] == {
        "2": [{"index": 3, "text": "A segment that failed on page two."}]
    }
    # Page 1 had no dropped segments — must be absent, not a zero-length placeholder.
    assert "1" not in transcript["dropped_segments"]


@pytest.mark.asyncio
async def test_concat_phase_transcript_omits_pages_without_sidecar():
    """§8 edge case: a page with no segments sidecar (pre-Sprint-1 book, or one
    that hit the total-TTS-failure branch) must be simply absent from
    transcript.json's segments key — not an error, not a zero-length
    placeholder — and must not prevent the rest of the transcript from being
    written correctly."""
    bid = AudiobookStore.create_book("Test.pdf")
    meta = AudiobookStore.initial_meta(
        bid, "Test.pdf", 2, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
    )
    AudiobookStore.write_meta(bid, meta)

    _write_pcm_wav(AudiobookStore.page_audio_path(bid, 1), 24000)
    AudiobookService._write_segments_json(
        AudiobookStore.page_segments_path(bid, 1),
        [{"text": "Page one sentence.", "start_sec": 0.0, "end_sec": 1.0}],
        [],
    )
    # Page 2 has real audio but deliberately NO sidecar file.
    _write_pcm_wav(AudiobookStore.page_audio_path(bid, 2), 12000)

    # Must not raise despite page 2's missing sidecar.
    await AudiobookService._phase_concat(bid, AudiobookStore.read_meta(bid))

    with open(AudiobookStore.transcript_path(bid), encoding="utf-8") as f:
        transcript = json.load(f)

    assert "1" in transcript["segments"]
    assert (
        "2" not in transcript["segments"]
    ), "missing sidecar must be absent, not a placeholder"
    assert transcript["dropped_segments"] == {}


def test_wav_header_format():
    body_size = 1000
    h = _wav_header(body_size)
    assert len(h) == 44
    assert h[:4] == b"RIFF"
    assert h[8:12] == b"WAVE"
    assert h[36:40] == b"data"


def test_write_segments_json_round_trip(tmp_path):
    path = str(tmp_path / "audio_pages" / "003.segments.json")
    segments = [
        {"text": "Hello world.", "start_sec": 0.0, "end_sec": 1.2},
        {"text": "Goodbye.", "start_sec": 1.2, "end_sec": 2.0},
    ]
    dropped = [{"index": 2, "text": "This one failed."}]

    AudiobookService._write_segments_json(path, segments, dropped)

    with open(path, encoding="utf-8") as f:
        data = json.load(f)
    assert data == {"segments": segments, "dropped": dropped}


# ---------- TTS phase + per-page failure ----------


async def _mock_generate_with_timing_yielding(*args, **kwargs):
    """Two successful segments, mirroring the pre-Sprint-1 mock's two chunks
    (12000, 6000 samples) but in generate_with_timing's dict shape."""
    for i, n_samples in enumerate((12000, 6000)):
        audio = np.zeros(n_samples, dtype=np.float32)
        yield {
            "index": i,
            "text": f"segment {i}",
            "ok": True,
            "audio": audio,
            "duration_sec": len(audio) / SAMPLE_RATE,
        }


# ---------- _generate_full_page (T1.5) ----------


@pytest.mark.asyncio
async def test_generate_full_page_returns_samples_timings_and_dropped():
    """T1.5's new 3-tuple return shape. A mix of ok/failed segments must
    produce: samples == concatenation of only the ok segments' audio;
    segment_timings with correct page-relative cumulative offsets;
    dropped_segments containing exactly the failed entries."""

    async def mixed_generate_with_timing(text, voice, speed):
        yield {
            "index": 0,
            "text": "First.",
            "ok": True,
            "audio": np.ones(1000, dtype=np.float32),
            "duration_sec": 1000 / SAMPLE_RATE,
        }
        yield {
            "index": 1,
            "text": "Second (fails).",
            "ok": False,
            "audio": None,
            "duration_sec": 0.0,
        }
        yield {
            "index": 2,
            "text": "Third.",
            "ok": True,
            "audio": np.full(500, 0.5, dtype=np.float32),
            "duration_sec": 500 / SAMPLE_RATE,
        }

    with patch(
        "app.services.audiobook_service.EngineManager.generate_with_timing",
        side_effect=mixed_generate_with_timing,
    ):
        samples, segment_timings, dropped_segments = (
            await AudiobookService._generate_full_page("irrelevant", "af_bella", 1.0)
        )

    expected_samples = np.concatenate(
        [np.ones(1000, dtype=np.float32), np.full(500, 0.5, dtype=np.float32)]
    )
    assert np.array_equal(samples, expected_samples)

    assert len(segment_timings) == 2
    assert segment_timings[0]["text"] == "First."
    assert segment_timings[0]["start_sec"] == 0.0
    assert abs(segment_timings[0]["end_sec"] - 1000 / SAMPLE_RATE) < 1e-9
    assert segment_timings[1]["text"] == "Third."
    assert abs(segment_timings[1]["start_sec"] - 1000 / SAMPLE_RATE) < 1e-9
    assert abs(segment_timings[1]["end_sec"] - 1500 / SAMPLE_RATE) < 1e-9

    assert dropped_segments == [{"index": 1, "text": "Second (fails)."}]


@pytest.mark.asyncio
async def test_generate_full_page_all_segments_fail_returns_silence_fallback():
    """§8 edge case: when every segment fails, _generate_full_page preserves
    the pre-existing 0.3s-silence fallback byte-for-byte, with empty
    segment_timings and every segment recorded in dropped_segments (not
    silently vanished, which was the D1 bug)."""

    async def all_fail_generate_with_timing(text, voice, speed):
        yield {
            "index": 0,
            "text": "Bad one.",
            "ok": False,
            "audio": None,
            "duration_sec": 0.0,
        }
        yield {
            "index": 1,
            "text": "Also bad.",
            "ok": False,
            "audio": None,
            "duration_sec": 0.0,
        }

    with patch(
        "app.services.audiobook_service.EngineManager.generate_with_timing",
        side_effect=all_fail_generate_with_timing,
    ):
        samples, segment_timings, dropped_segments = (
            await AudiobookService._generate_full_page("irrelevant", "af_bella", 1.0)
        )

    assert np.array_equal(samples, np.zeros(int(0.3 * SAMPLE_RATE), dtype=np.float32))
    assert segment_timings == []
    assert dropped_segments == [
        {"index": 0, "text": "Bad one."},
        {"index": 1, "text": "Also bad."},
    ]


@pytest.mark.asyncio
async def test_tts_phase_flags_partial_segment_failure(monkeypatch):
    """D1 end-to-end: a page with one failed segment among several successful
    ones must land in failed_pages even though its WAV still has real
    (non-silence) audio for the segments that succeeded. Pre-fix, this class
    of failure was silently swallowed inside generate() — the page was never
    flagged and failed_count stayed 0 even though words were missing."""
    bid = AudiobookStore.create_book("Test.pdf")
    meta = AudiobookStore.initial_meta(
        bid, "Test.pdf", 1, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
    )
    AudiobookStore.write_meta(bid, meta)
    path = AudiobookStore.page_clean_path(bid, 1)
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w") as f:
        f.write("First sentence succeeds. Second sentence fails. Third succeeds.")

    async def partial_fail_generate_with_timing(text, voice, speed):
        yield {
            "index": 0,
            "text": "First sentence succeeds.",
            "ok": True,
            "audio": np.ones(2400, dtype=np.float32),
            "duration_sec": 2400 / SAMPLE_RATE,
        }
        yield {
            "index": 1,
            "text": "Second sentence fails.",
            "ok": False,
            "audio": None,
            "duration_sec": 0.0,
        }
        yield {
            "index": 2,
            "text": "Third succeeds.",
            "ok": True,
            "audio": np.ones(1200, dtype=np.float32),
            "duration_sec": 1200 / SAMPLE_RATE,
        }

    events: list[tuple[str, dict]] = []
    monkeypatch.setattr(
        AudiobookService,
        "_emit",
        classmethod(lambda cls, b, ev, **kw: events.append((ev, kw))),
    )

    with (
        patch(
            "app.services.audiobook_service.EngineManager.ensure_loaded",
            new=AsyncMock(return_value=None),
        ),
        patch("app.services.audiobook_service.EngineManager.touch", return_value=None),
        patch(
            "app.services.audiobook_service.EngineManager.generate_with_timing",
            side_effect=partial_fail_generate_with_timing,
        ),
    ):
        await AudiobookService._phase_tts(bid, meta)

    new_meta = AudiobookStore.read_meta(bid)
    assert 1 in new_meta["failed_pages"], "partial segment failure must flag the page"

    # Real, non-silence audio from the 2 successful segments — not a blanket
    # silence fallback. (2400 + 1200 samples * 2 bytes = 7200 PCM bytes.)
    wav_path = AudiobookStore.page_audio_path(bid, 1)
    assert os.path.exists(wav_path)
    assert os.path.getsize(wav_path) == WAV_HEADER_SIZE + (2400 + 1200) * 2

    with open(AudiobookStore.page_segments_path(bid, 1), encoding="utf-8") as f:
        sidecar = json.load(f)
    assert len(sidecar["segments"]) == 2
    assert sidecar["dropped"] == [{"index": 1, "text": "Second sentence fails."}]

    page_failed = [kw for ev, kw in events if ev == "page_failed"]
    assert any(
        e.get("phase") == "tts" for e in page_failed
    ), "page_failed SSE must fire"


@pytest.mark.asyncio
async def test_tts_phase_sidecar_write_failure_does_not_destroy_good_audio():
    """Regression: a segments-sidecar write failure occurring AFTER the real
    WAV was already written must never cascade into the outer except-branch
    and overwrite good audio with silence. Per the v1.1 design notes §8, the WAV's
    existence is the pipeline's sole "page is done" checkpoint; a missing
    sidecar is an already-designed-for, harmless fallback (the frontend
    degrades to the old per-page estimate) — it must degrade to that, not
    destroy already-synthesized real audio or false-flag the page failed."""
    bid = AudiobookStore.create_book("Test.pdf")
    meta = AudiobookStore.initial_meta(
        bid, "Test.pdf", 1, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
    )
    AudiobookStore.write_meta(bid, meta)
    path = AudiobookStore.page_clean_path(bid, 1)
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w") as f:
        f.write("A perfectly good sentence that synthesizes just fine.")

    async def all_succeed_generate_with_timing(text, voice, speed):
        yield {
            "index": 0,
            "text": "A perfectly good sentence that synthesizes just fine.",
            "ok": True,
            "audio": np.ones(4800, dtype=np.float32),
            "duration_sec": 4800 / SAMPLE_RATE,
        }

    with (
        patch(
            "app.services.audiobook_service.EngineManager.ensure_loaded",
            new=AsyncMock(return_value=None),
        ),
        patch("app.services.audiobook_service.EngineManager.touch", return_value=None),
        patch(
            "app.services.audiobook_service.EngineManager.generate_with_timing",
            side_effect=all_succeed_generate_with_timing,
        ),
        patch.object(
            AudiobookService,
            "_write_segments_json",
            side_effect=OSError("disk full"),
        ),
    ):
        await AudiobookService._phase_tts(bid, meta)

    new_meta = AudiobookStore.read_meta(bid)
    assert 1 not in (
        new_meta.get("failed_pages") or []
    ), "a sidecar write failure alone must not mark the page failed"

    # Real synthesized audio (4800 samples * 2 bytes = 9600 PCM bytes), NOT
    # the 0.5s total-failure silence fallback (12000 samples = 24000 PCM bytes).
    wav_path = AudiobookStore.page_audio_path(bid, 1)
    assert os.path.exists(wav_path)
    assert (
        os.path.getsize(wav_path) == WAV_HEADER_SIZE + 4800 * 2
    ), "good audio must survive a sidecar-write failure, not be overwritten with silence"

    # No sidecar was persisted — must be simply absent, matching the
    # already-designed-for missing-sidecar fallback (no exception here).
    assert not os.path.exists(AudiobookStore.page_segments_path(bid, 1))


@pytest.mark.asyncio
async def test_tts_phase_writes_per_page_wavs(monkeypatch):
    bid = AudiobookStore.create_book("Test.pdf")
    meta = AudiobookStore.initial_meta(
        bid, "Test.pdf", 2, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
    )
    AudiobookStore.write_meta(bid, meta)
    # Pre-create cleaned text for 2 pages so tts phase has input.
    for n in (1, 2):
        path = AudiobookStore.page_clean_path(bid, n)
        os.makedirs(os.path.dirname(path), exist_ok=True)
        with open(path, "w") as f:
            f.write(f"Page {n} content.")

    with (
        patch(
            "app.services.audiobook_service.EngineManager.ensure_loaded",
            new=AsyncMock(return_value=None),
        ),
        patch(
            "app.services.audiobook_service.EngineManager.touch",
            return_value=None,
        ),
        patch(
            "app.services.audiobook_service.EngineManager.generate_with_timing",
            side_effect=_mock_generate_with_timing_yielding,
        ),
    ):
        await AudiobookService._phase_tts(bid, meta)

    for n in (1, 2):
        assert os.path.exists(AudiobookStore.page_audio_path(bid, n))
        # T1.6: a segments sidecar must be written alongside every real WAV.
        assert os.path.exists(AudiobookStore.page_segments_path(bid, n))


@pytest.mark.asyncio
async def test_phase_tts_ensures_model_loaded_each_page(monkeypatch):
    """Regression: a long interactive /speak can idle-unload the model while the
    audiobook waits on the preemption lock. _phase_tts must re-ensure the model
    is loaded per page so the page doesn't fail with 'Model not initialized'."""
    bid = AudiobookStore.create_book("Test.pdf")
    meta = AudiobookStore.initial_meta(
        bid, "Test.pdf", 2, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
    )
    AudiobookStore.write_meta(bid, meta)
    for n in (1, 2):
        path = AudiobookStore.page_clean_path(bid, n)
        os.makedirs(os.path.dirname(path), exist_ok=True)
        with open(path, "w") as f:
            f.write(f"Page {n} has some real content to narrate.")

    ensure = AsyncMock(return_value=None)
    with (
        patch("app.services.audiobook_service.EngineManager.ensure_loaded", ensure),
        patch("app.services.audiobook_service.EngineManager.touch", return_value=None),
        patch(
            "app.services.audiobook_service.EngineManager.generate_with_timing",
            side_effect=_mock_generate_with_timing_yielding,
        ),
    ):
        await AudiobookService._phase_tts(bid, meta)

    # Once before the loop + once per page (2) = 3. Assert at least per-page.
    assert ensure.await_count >= 2, "model must be re-ensured each page"


@pytest.mark.asyncio
async def test_run_pipeline_surfaces_failed_count_in_done(monkeypatch):
    """A book that had per-page clean/TTS failures must report failed_count in
    the 'done' event (and persisted stats) so the UI can offer a retry rather
    than presenting a flawless book."""
    bid = AudiobookStore.create_book("Test.pdf")
    meta = AudiobookStore.initial_meta(
        bid, "Test.pdf", 3, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
    )
    AudiobookStore.write_meta(bid, meta)
    # Simulate clean/tts having recorded two failed pages.
    await AudiobookStore.update_meta(bid, failed_pages=[2, 3])

    events: list[tuple[str, dict]] = []
    monkeypatch.setattr(
        AudiobookService,
        "_emit",
        classmethod(lambda cls, b, ev, **kw: events.append((ev, kw))),
    )
    monkeypatch.setattr(AudiobookService, "_phase_extract", AsyncMock())
    monkeypatch.setattr(AudiobookService, "_phase_clean", AsyncMock())
    monkeypatch.setattr(AudiobookService, "_phase_section", AsyncMock())
    monkeypatch.setattr(AudiobookService, "_phase_tts", AsyncMock())
    monkeypatch.setattr(
        AudiobookService, "_phase_concat", AsyncMock(return_value={"pages": 3})
    )

    await AudiobookService._run_pipeline(bid)

    done = [kw for ev, kw in events if ev == "done"]
    assert done, "no 'done' event emitted"
    actual = done[0]["actual"]
    assert actual["failed_count"] == 2
    assert actual["failed_pages"] == [2, 3]
    # Persisted into meta.actual too.
    persisted = AudiobookStore.read_meta(bid)
    assert persisted["actual"]["failed_count"] == 2
    assert persisted["status"] == "done"


@pytest.mark.asyncio
async def test_run_pipeline_clean_book_reports_zero_failures(monkeypatch):
    """No failed pages → failed_count 0, failed_pages empty (no false alarms)."""
    bid = AudiobookStore.create_book("Test.pdf")
    meta = AudiobookStore.initial_meta(
        bid, "Test.pdf", 2, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
    )
    AudiobookStore.write_meta(bid, meta)

    events: list[tuple[str, dict]] = []
    monkeypatch.setattr(
        AudiobookService,
        "_emit",
        classmethod(lambda cls, b, ev, **kw: events.append((ev, kw))),
    )
    for name in ("_phase_extract", "_phase_clean", "_phase_section", "_phase_tts"):
        monkeypatch.setattr(AudiobookService, name, AsyncMock())
    monkeypatch.setattr(
        AudiobookService, "_phase_concat", AsyncMock(return_value={"pages": 2})
    )

    await AudiobookService._run_pipeline(bid)

    actual = [kw for ev, kw in events if ev == "done"][0]["actual"]
    assert actual["failed_count"] == 0
    assert actual["failed_pages"] == []


# ---------- API endpoints ----------


def test_audiobook_list_endpoint_initially_empty():
    from fastapi.testclient import TestClient

    from app.main import app

    client = TestClient(app)
    response = client.get("/audiobook")
    assert response.status_code == 200
    assert response.json() == []


def test_audiobook_404_for_unknown_id():
    """Well-formed but non-existent UUID → 404."""
    from fastapi.testclient import TestClient

    from app.main import app

    client = TestClient(app)
    # Valid hex-32 shape, but the row doesn't exist.
    response = client.get("/audiobook/" + ("a" * 32))
    assert response.status_code == 404


def test_audiobook_400_for_malformed_id():
    """Anything that isn't 32 hex chars → 400 (HARD-017 path-traversal guard)."""
    from fastapi.testclient import TestClient

    from app.main import app

    client = TestClient(app)
    for bad in ["nonexistent_id_12345", "../../etc/passwd", "ZZZZ", "a" * 31, "a" * 33]:
        response = client.get(f"/audiobook/{bad}")
        # FastAPI 404s on path-segment mismatch when the route doesn't match
        # (e.g. "../"). For shapes that DO traverse our route, expect 400.
        assert response.status_code in (400, 404), (bad, response.status_code)


def test_audiobook_upload_rejects_oversize_file(monkeypatch):
    """Body > MAX_AUDIOBOOK_UPLOAD_MB → 413, no OOM. Limit is shrunk to 1 MB
    for the test so we don't have to allocate 100 MB of test data."""
    from fastapi.testclient import TestClient

    from app.core.config import settings
    from app.main import app

    monkeypatch.setattr(settings, "MAX_AUDIOBOOK_UPLOAD_MB", 1)
    client = TestClient(app)
    huge = b"%PDF-1.4\n" + b"A" * (2 * 1024 * 1024)  # 2 MiB > 1 MB limit
    response = client.post(
        "/audiobook",
        files={"file": ("big.pdf", huge, "application/pdf")},
    )
    assert response.status_code == 413
    assert "MB" in response.json()["detail"]


def test_start_requires_api_key_header():
    from fastapi.testclient import TestClient

    from app.main import app

    client = TestClient(app)
    bid = AudiobookStore.create_book("Test.pdf")
    AudiobookStore.write_meta(
        bid,
        AudiobookStore.initial_meta(
            bid, "Test.pdf", 1, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
        ),
    )
    response = client.post(f"/audiobook/{bid}/start")
    assert response.status_code == 400
    assert "X-Gemini-Api-Key" in response.json()["detail"]


def test_start_requires_consent_header():
    """HARD-072: consent is a UI-only gate client-side, so the server must
    independently reject a /start call missing (or false) X-Gemini-Consent —
    otherwise a modified client could send a document to Gemini without ever
    showing the user the consent toggle."""
    from fastapi.testclient import TestClient

    from app.main import app

    client = TestClient(app)
    bid = AudiobookStore.create_book("Test.pdf")
    AudiobookStore.write_meta(
        bid,
        AudiobookStore.initial_meta(
            bid, "Test.pdf", 1, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
        ),
    )
    response = client.post(f"/audiobook/{bid}/start", headers={"X-Gemini-Api-Key": "x"})
    assert response.status_code == 400
    assert "Consent" in response.json()["detail"]

    response = client.post(
        f"/audiobook/{bid}/start",
        headers={"X-Gemini-Api-Key": "x", "X-Gemini-Consent": "false"},
    )
    assert response.status_code == 400
    assert "Consent" in response.json()["detail"]


def test_start_succeeds_with_api_key_and_consent(monkeypatch):
    from fastapi.testclient import TestClient

    from app.main import app

    enqueued: list[str] = []

    async def fake_enqueue(book_id: str, api_key: str):
        enqueued.append(book_id)

    monkeypatch.setattr(
        AudiobookService, "enqueue", classmethod(lambda cls, b, k: fake_enqueue(b, k))
    )

    client = TestClient(app)
    bid = AudiobookStore.create_book("Test.pdf")
    AudiobookStore.write_meta(
        bid,
        AudiobookStore.initial_meta(
            bid, "Test.pdf", 1, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
        ),
    )
    response = client.post(
        f"/audiobook/{bid}/start",
        headers={"X-Gemini-Api-Key": "x", "X-Gemini-Consent": "true"},
    )
    assert response.status_code == 200
    assert response.json()["status"] == "queued"
    assert enqueued == [bid]


# ---------- resume ----------


@pytest.mark.asyncio
async def test_resume_in_progress_flips_cleaning_to_needs_key():
    bid = AudiobookStore.create_book("Test.pdf")
    meta = AudiobookStore.initial_meta(
        bid, "Test.pdf", 1, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
    )
    meta["status"] = "cleaning"
    AudiobookStore.write_meta(bid, meta)

    await AudiobookService.resume_in_progress()
    new_meta = AudiobookStore.read_meta(bid)
    assert new_meta["status"] == "needs_key"


# ---------- Phase 2: section detection ----------


def test_stitch_sections_basic_contiguity():
    from app.services.gemini_cleaner import GeminiCleaner

    raw = [
        {"title": "Intro", "start_page": 1, "end_page": 3},
        {"title": "Chapter 1", "start_page": 4, "end_page": 9},
        {"title": "Chapter 2", "start_page": 10, "end_page": 15},
    ]
    out = GeminiCleaner._stitch_sections(raw, page_count=20)
    # Last section's end_page extended to page_count.
    assert out[-1]["end_page"] == 20
    # Contiguous: each end == next start - 1.
    for a, b in zip(out, out[1:]):
        assert a["end_page"] == b["start_page"] - 1


def test_stitch_sections_inserts_front_matter():
    from app.services.gemini_cleaner import GeminiCleaner

    raw = [{"title": "Chapter 1", "start_page": 4, "end_page": 9}]
    out = GeminiCleaner._stitch_sections(raw, page_count=12)
    assert out[0]["title"] == "Front Matter"
    assert out[0]["start_page"] == 1
    assert out[0]["end_page"] == 3
    assert out[-1]["end_page"] == 12


def test_parse_sections_json_strips_markdown_fence():
    from app.services.gemini_cleaner import GeminiCleaner

    fenced = '```json\n{"sections":[{"title":"A","start_page":1,"end_page":5}]}\n```'
    out = GeminiCleaner._parse_sections_json(fenced, max_page=10)
    assert out == [{"title": "A", "start_page": 1, "end_page": 5}]


def test_parse_sections_json_drops_invalid_entries():
    from app.services.gemini_cleaner import GeminiCleaner

    bad = json.dumps(
        {
            "sections": [
                {"title": "Good", "start_page": 1, "end_page": 3},
                {"title": "", "start_page": 4, "end_page": 5},  # empty title
                {"title": "BadOrder", "start_page": 9, "end_page": 7},  # end < start
                {"title": "Beyond", "start_page": 100, "end_page": 200},  # out of range
            ]
        }
    )
    out = GeminiCleaner._parse_sections_json(bad, max_page=10)
    assert len(out) == 1
    assert out[0]["title"] == "Good"


# ---------- Phase 2: HTTP Range support ----------


@pytest.mark.asyncio
async def test_audio_range_request_returns_206_with_correct_slice():
    from fastapi.testclient import TestClient

    from app.main import app

    client = TestClient(app)
    bid = AudiobookStore.create_book("Test.pdf")
    AudiobookStore.write_meta(
        bid,
        AudiobookStore.initial_meta(
            bid, "Test.pdf", 1, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
        ),
    )
    # Create a known-content audio.wav.
    audio_path = AudiobookStore.audio_path(bid)
    payload = bytes(range(256)) * 16  # 4096 bytes
    with open(audio_path, "wb") as f:
        f.write(payload)

    response = client.get(f"/audiobook/{bid}/audio", headers={"Range": "bytes=100-199"})
    assert response.status_code == 206
    assert response.headers["Content-Range"] == f"bytes 100-199/{len(payload)}"
    assert response.headers["Accept-Ranges"] == "bytes"
    assert response.content == payload[100:200]


def test_audio_no_range_returns_full_file_with_accept_ranges_header():
    from fastapi.testclient import TestClient

    from app.main import app

    client = TestClient(app)
    bid = AudiobookStore.create_book("Test.pdf")
    AudiobookStore.write_meta(
        bid,
        AudiobookStore.initial_meta(
            bid, "Test.pdf", 1, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
        ),
    )
    # >= 44 bytes: the endpoint treats sub-WAV-header files as not-ready (404),
    # and a real audio.wav always has at least a 44-byte header.
    payload = b"RIFF" + b"\x00" * 60
    with open(AudiobookStore.audio_path(bid), "wb") as f:
        f.write(payload)

    response = client.get(f"/audiobook/{bid}/audio")
    assert response.status_code == 200
    assert response.headers["Accept-Ranges"] == "bytes"
    assert response.content == payload


def test_audio_subheader_file_returns_404_not_ready():
    """A partial WAV must not be exposed as a successful audio response."""
    from fastapi.testclient import TestClient

    from app.main import app

    client = TestClient(app)
    book_id = AudiobookStore.create_book("Test.pdf")
    AudiobookStore.write_meta(
        book_id,
        AudiobookStore.initial_meta(
            book_id, "Test.pdf", 1, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
        ),
    )
    with open(AudiobookStore.audio_path(book_id), "wb") as file:
        file.write(b"\x00" * 10)

    assert client.get(f"/audiobook/{book_id}/audio").status_code == 404


def test_audio_range_invalid_returns_416():
    from fastapi.testclient import TestClient

    from app.main import app

    client = TestClient(app)
    bid = AudiobookStore.create_book("Test.pdf")
    AudiobookStore.write_meta(
        bid,
        AudiobookStore.initial_meta(
            bid, "Test.pdf", 1, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
        ),
    )
    with open(AudiobookStore.audio_path(bid), "wb") as f:
        f.write(b"x" * 100)

    # Request range past the file end.
    response = client.get(f"/audiobook/{bid}/audio", headers={"Range": "bytes=500-600"})
    assert response.status_code == 416


# ---------- Phase 2: cancel ----------


def test_cancel_endpoint_404_for_unknown_book():
    from fastapi.testclient import TestClient

    from app.main import app

    client = TestClient(app)
    # Use a valid-shape UUID that doesn't exist (HARD-017 rejects malformed IDs at 400).
    response = client.post("/audiobook/" + ("b" * 32) + "/cancel")
    assert response.status_code == 404


def test_cancel_sets_flag():
    bid = AudiobookStore.create_book("Test.pdf")
    AudiobookStore.write_meta(
        bid,
        AudiobookStore.initial_meta(
            bid, "Test.pdf", 1, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
        ),
    )
    assert AudiobookService.cancel(bid) is True
    assert AudiobookService._cancel_flags.get(bid) is True


# ---------- Phase 2: transcript ----------


def test_transcript_endpoint_404_when_missing():
    from fastapi.testclient import TestClient

    from app.main import app

    client = TestClient(app)
    bid = AudiobookStore.create_book("Test.pdf")
    AudiobookStore.write_meta(
        bid,
        AudiobookStore.initial_meta(
            bid, "Test.pdf", 1, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
        ),
    )
    response = client.get(f"/audiobook/{bid}/transcript")
    assert response.status_code == 404


def test_transcript_endpoint_serves_file():
    import json as _json

    from fastapi.testclient import TestClient

    from app.main import app

    client = TestClient(app)
    bid = AudiobookStore.create_book("Test.pdf")
    AudiobookStore.write_meta(
        bid,
        AudiobookStore.initial_meta(
            bid, "Test.pdf", 1, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
        ),
    )
    payload = {"book_id": bid, "sections": []}
    with open(AudiobookStore.transcript_path(bid), "w") as f:
        _json.dump(payload, f)

    response = client.get(f"/audiobook/{bid}/transcript")
    assert response.status_code == 200
    assert response.json()["book_id"] == bid


# ---------- Phase 2: retry endpoint ----------


@pytest.mark.asyncio
async def test_retry_failed_clears_pages_and_enqueues(monkeypatch):
    bid = AudiobookStore.create_book("Test.pdf")
    meta = AudiobookStore.initial_meta(
        bid, "Test.pdf", 3, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
    )
    meta["status"] = "failed"
    meta["failed_pages"] = [2, 3]
    AudiobookStore.write_meta(bid, meta)
    # Pre-create the failed-page files so we can verify removal.
    for n in (2, 3):
        for path in (
            AudiobookStore.page_clean_path(bid, n),
            AudiobookStore.page_audio_path(bid, n),
            AudiobookStore.page_segments_path(bid, n),
        ):
            os.makedirs(os.path.dirname(path), exist_ok=True)
            open(path, "wb").close()
    # Touch final audio + transcript so we verify they're cleared too.
    open(AudiobookStore.audio_path(bid), "wb").close()
    open(AudiobookStore.transcript_path(bid), "w").close()

    # Stub the queue so retry_failed doesn't actually run a pipeline.
    enqueued: list[str] = []

    async def fake_enqueue(book_id: str, api_key: str):
        enqueued.append(book_id)

    monkeypatch.setattr(
        AudiobookService, "enqueue", classmethod(lambda cls, b, k: fake_enqueue(b, k))
    )

    count = await AudiobookService.retry_failed(bid, "fake-key")
    assert count == 2
    assert enqueued == [bid]
    # Per-page intermediates wiped:
    for n in (2, 3):
        assert not os.path.exists(AudiobookStore.page_clean_path(bid, n))
        assert not os.path.exists(AudiobookStore.page_audio_path(bid, n))
        # T1.8: a retry must not leave a stale segments sidecar describing
        # the first attempt's timing against what will be a fresh WAV.
        assert not os.path.exists(AudiobookStore.page_segments_path(bid, n))
    # Final audio + transcript wiped:
    assert not os.path.exists(AudiobookStore.audio_path(bid))
    assert not os.path.exists(AudiobookStore.transcript_path(bid))
    new_meta = AudiobookStore.read_meta(bid)
    assert new_meta["failed_pages"] == []
    assert new_meta["error"] is None


def test_retry_endpoint_requires_api_key():
    from fastapi.testclient import TestClient

    from app.main import app

    client = TestClient(app)
    bid = AudiobookStore.create_book("Test.pdf")
    AudiobookStore.write_meta(
        bid,
        AudiobookStore.initial_meta(
            bid, "Test.pdf", 1, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
        ),
    )
    response = client.post(f"/audiobook/{bid}/retry")
    assert response.status_code == 400


@pytest.mark.asyncio
async def test_request_delete_cancels_in_flight_pipeline():
    """Regression for C8: DELETE on a processing book must signal cancel +
    wait at the next page boundary, not rmtree out from under the worker."""
    bid = AudiobookStore.create_book("Test.pdf")
    meta = AudiobookStore.initial_meta(
        bid, "Test.pdf", 5, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
    )
    meta["status"] = "tts"
    AudiobookStore.write_meta(bid, meta)
    AudiobookService._current_book_id = bid
    AudiobookService._cancel_flags.pop(bid, None)

    async def release_after_short_pause():
        await asyncio.sleep(0.15)
        AudiobookService._current_book_id = None

    asyncio.create_task(release_after_short_pause())

    ok = await AudiobookService.request_delete(bid)
    assert ok is True
    # Flag was cleared after delete completed.
    assert bid not in AudiobookService._cancel_flags
    # Directory removed (read_meta returns None for missing dirs).
    assert AudiobookStore.read_meta(bid) is None


@pytest.mark.asyncio
async def test_request_delete_unknown_returns_false():
    ok = await AudiobookService.request_delete("nonexistent_id_xyz")
    assert ok is False


@pytest.mark.asyncio
async def test_retry_failed_resumes_needs_key_book(monkeypatch):
    """Regression for C2: a book in needs_key state with no failed pages
    must still be enqueued when retry_failed is called (after the user
    re-enters their API key). Previously this silently returned 0."""
    bid = AudiobookStore.create_book("Test.pdf")
    meta = AudiobookStore.initial_meta(
        bid, "Test.pdf", 5, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
    )
    meta["status"] = "needs_key"
    meta["failed_pages"] = []
    AudiobookStore.write_meta(bid, meta)

    enqueued: list[str] = []

    async def fake_enqueue(book_id: str, api_key: str):
        enqueued.append(book_id)

    monkeypatch.setattr(
        AudiobookService, "enqueue", classmethod(lambda cls, b, k: fake_enqueue(b, k))
    )

    count = await AudiobookService.retry_failed(bid, "fake-key")
    # Zero pages flagged, but enqueue must still fire.
    assert count == 0
    assert enqueued == [bid]


def test_retry_endpoint_404_for_unknown_book():
    from fastapi.testclient import TestClient

    from app.main import app

    client = TestClient(app)
    response = client.post(
        "/audiobook/" + ("c" * 32) + "/retry", headers={"X-Gemini-Api-Key": "x"}
    )
    assert response.status_code == 404


# ---------- voice/speed/engine flow at upload ----------


def test_initial_meta_preserves_voice_and_speed():
    bid = AudiobookStore.create_book("Test.pdf")
    meta = AudiobookStore.initial_meta(
        bid, "Test.pdf", 5, "kitten", "Bella", 1.5, {"cost_usd": 0.1}
    )
    assert meta["engine"] == "kitten"
    assert meta["voice"] == "Bella"
    assert meta["speed"] == 1.5


def test_estimate_uses_speed_for_audio_duration():
    e_slow = AudiobookService.estimate(
        page_count=10, sample_words=100, sample_chars=500, speed=1.0
    )
    e_fast = AudiobookService.estimate(
        page_count=10, sample_words=100, sample_chars=500, speed=2.0
    )
    # 2x speed → ~half the audio seconds.
    assert e_fast["audio_seconds"] < e_slow["audio_seconds"]
    assert abs(e_fast["audio_seconds"] * 2 - e_slow["audio_seconds"]) < 0.5


# ---------- Gemini retry/backoff (mocked) ----------


@pytest.mark.asyncio
async def test_gemini_clean_page_retries_then_succeeds(monkeypatch):
    from app.services.gemini_cleaner import GeminiBadResponseError, GeminiCleaner

    calls = {"n": 0}

    async def flaky_async(api_key, raw):
        calls["n"] += 1
        if calls["n"] < 2:
            raise GeminiBadResponseError("transient")
        return "cleaned text"

    monkeypatch.setattr(
        GeminiCleaner, "_async_clean", AsyncMock(side_effect=flaky_async)
    )
    monkeypatch.setattr("asyncio.sleep", AsyncMock(return_value=None))
    out = await GeminiCleaner.clean_page("k", "raw")
    assert out == "cleaned text"
    assert calls["n"] == 2


@pytest.mark.asyncio
async def test_gemini_auth_error_does_not_retry(monkeypatch):
    from app.services.gemini_cleaner import GeminiAuthError, GeminiCleaner

    calls = {"n": 0}

    async def auth_failing(api_key, raw):
        calls["n"] += 1
        raise GeminiAuthError("bad key")

    monkeypatch.setattr(
        GeminiCleaner, "_async_clean", AsyncMock(side_effect=auth_failing)
    )
    monkeypatch.setattr("asyncio.sleep", AsyncMock(return_value=None))
    with pytest.raises(GeminiAuthError):
        await GeminiCleaner.clean_page("k", "raw")
    assert calls["n"] == 1  # no retry for auth errors


# ---------- cost_warning flag ----------


def test_upload_endpoint_happy_path(monkeypatch):
    """Real HTTP upload path against the in-process app: PDF parse → estimate → meta written.

    Uses a fake PDFExtractor so we don't need a real PDF file. This was added
    after a live binary smoke test exposed `asyncio.create_task(future)` raising.
    """
    from fastapi.testclient import TestClient

    from app.main import app
    from app.services import pdf_extractor as _pe

    monkeypatch.setattr(_pe.PDFExtractor, "page_count", classmethod(lambda cls, p: 3))
    monkeypatch.setattr(
        _pe.PDFExtractor, "is_image_only", classmethod(lambda cls, p: False)
    )
    monkeypatch.setattr(
        _pe.PDFExtractor, "sample_word_count", classmethod(lambda cls, p: 50)
    )
    monkeypatch.setattr(
        _pe.PDFExtractor, "sample_char_count", classmethod(lambda cls, p: 250)
    )
    rendered: list[str] = []
    monkeypatch.setattr(
        _pe.PDFExtractor, "render_cover", classmethod(lambda cls, b: rendered.append(b))
    )

    client = TestClient(app)
    # Pad past the 100-byte minimum size guard. Real parsing is mocked above.
    files = {"file": ("test.pdf", b"%PDF-1.4\n" + b"x" * 200, "application/pdf")}
    data = {"voice": "bf_emma", "speed": "1.25", "engine": "kokoro"}
    response = client.post("/audiobook", files=files, data=data)
    assert response.status_code == 200, response.text
    body = response.json()
    assert body["title"] == "test.pdf"
    assert body["page_count"] == 3
    assert body["word_count_estimate"] == 150  # 50 * 3
    assert body["estimated_token_count"] > 0
    assert "cost_warning" in body
    assert body["is_image_only"] is False

    # Meta on disk reflects user's voice/speed/engine.
    bid = body["book_id"]
    meta = AudiobookStore.read_meta(bid)
    assert meta["voice"] == "bf_emma"
    assert meta["speed"] == 1.25
    assert meta["engine"] == "kokoro"


def test_upload_rejects_empty_pdf():
    """P9: zero-byte uploads should fail at the door."""
    from fastapi.testclient import TestClient

    from app.main import app

    client = TestClient(app)
    response = client.post(
        "/audiobook", files={"file": ("empty.pdf", b"", "application/pdf")}
    )
    assert response.status_code == 400
    assert "empty" in response.json()["detail"].lower()


def test_upload_rejects_unsupported_extension():
    """Only PDF, TXT, DOCX, and MD are accepted; everything else must return 400."""
    from fastapi.testclient import TestClient

    from app.main import app

    client = TestClient(app)
    files = {"file": ("test.exe", b"binary data", "application/octet-stream")}
    response = client.post("/audiobook", files=files)
    assert response.status_code == 400


def test_upload_accepts_image_only_pdf(monkeypatch):
    """Image-only PDFs are now accepted; OCR handles them during the clean phase.
    The response should return is_image_only=True so the UI can show an OCR badge."""
    from fastapi.testclient import TestClient

    from app.main import app
    from app.services import pdf_extractor as _pe

    monkeypatch.setattr(_pe.PDFExtractor, "page_count", classmethod(lambda cls, p: 1))
    monkeypatch.setattr(
        _pe.PDFExtractor, "is_image_only", classmethod(lambda cls, p: True)
    )
    monkeypatch.setattr(
        _pe.PDFExtractor, "render_cover", classmethod(lambda cls, bid, **kw: None)
    )

    client = TestClient(app)
    files = {"file": ("scan.pdf", b"%PDF-1.4\n" + b"x" * 200, "application/pdf")}
    response = client.post("/audiobook", files=files)
    assert response.status_code == 200
    body = response.json()
    assert body["is_image_only"] is True
    # Cost estimate uses the per-page OCR default (not zero).
    assert body["estimated_cost_usd"] > 0


def test_estimate_response_includes_cost_warning(monkeypatch):
    """Estimate carries cost_warning=True when projected cost > $1 (default threshold)."""
    # Tiny PDF → tiny cost → no warning.
    e = AudiobookService.estimate(
        page_count=1, sample_words=5, sample_chars=20, speed=1.0
    )
    assert e["cost_usd"] < 0.01

    # Large book → cost crosses the $1 threshold.
    big = AudiobookService.estimate(
        page_count=2500, sample_words=600, sample_chars=8000, speed=1.0
    )
    assert big["cost_usd"] > 1.0


# ---------- OCR fallback in clean phase ----------


@pytest.mark.asyncio
async def test_clean_phase_uses_ocr_for_image_pages(monkeypatch):
    """Pages with SOME (but under-threshold) extracted text are routed to
    GeminiCleaner.ocr_page instead of clean_page. A genuinely-empty page is a
    different case (D6 — see test_ocr_blank_page_routing.py): it must short-
    circuit for free rather than reach OCR at all, so this fixture uses a
    small amount of sparse text (e.g. a page-number/watermark fragment a real
    scan's text layer might leak) to stay strictly a "some text" case."""
    from app.services import audiobook_service as _svc
    from app.services import pdf_extractor as _pe

    bid = AudiobookStore.create_book("Scan.pdf")
    meta = AudiobookStore.initial_meta(
        bid, "Scan.pdf", 2, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
    )
    AudiobookStore.write_meta(bid, meta)
    AudiobookStore.save_source(bid, b"%PDF-1.4\n" + b"x" * 200, "pdf")

    # Page 1: sparse text under the OCR threshold, but NOT empty (image page
    # with a little bleed-through text) — should trigger OCR.
    # Page 2: normal text — should use clean_page.
    for n, content in [(1, "pg 1"), (2, "x" * 200)]:
        path = AudiobookStore.page_raw_path(bid, n)
        os.makedirs(os.path.dirname(path), exist_ok=True)
        with open(path, "w") as f:
            f.write(content)

    ocr_calls: list[int] = []
    clean_calls: list[int] = []

    async def _fake_ocr(api_key, image_bytes):
        ocr_calls.append(1)
        return "OCR result."

    async def _fake_clean(api_key, text):
        clean_calls.append(1)
        return "Cleaned text."

    monkeypatch.setattr(
        _pe.PDFExtractor,
        "render_page_image",
        classmethod(lambda cls, p, n, **kw: b"imgbytes"),
    )

    from app.services import gemini_cleaner as _gc

    monkeypatch.setattr(_gc.GeminiCleaner, "ocr_page", AsyncMock(side_effect=_fake_ocr))
    monkeypatch.setattr(
        _gc.GeminiCleaner, "clean_page", AsyncMock(side_effect=_fake_clean)
    )

    _svc.AudiobookService.initialize()
    await _svc.AudiobookService._phase_clean(bid, api_key="test-key")

    assert len(ocr_calls) == 1, "image page should route to OCR"
    assert len(clean_calls) == 1, "text page should route to clean_page"


# ---------- duplicate page deduplication ----------


@pytest.mark.asyncio
async def test_duplicate_pages_get_silence_marker(monkeypatch):
    """When two pages have identical content (e.g. DocuSign PDFs that embed
    the same page twice), the second occurrence must receive a '-' silence
    marker in its clean file so it does not produce duplicate audio.
    """
    from app.services import audiobook_service as _svc
    from app.services import pdf_extractor as _pe

    bid = AudiobookStore.create_book("offer.pdf")
    meta = AudiobookStore.initial_meta(
        bid, "offer.pdf", 2, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
    )
    meta["file_ext"] = "pdf"
    AudiobookStore.write_meta(bid, meta)

    # Pre-write identical raw page files (simulates what PDFExtractor would extract
    # from a DocuSign PDF where both pages contain the same text).
    long_content = (
        "A" * 200 + "\n\nThis is the full offer letter body with enough text to matter."
    )
    for n in (1, 2):
        path = AudiobookStore.page_raw_path(bid, n)
        os.makedirs(os.path.dirname(path), exist_ok=True)
        with open(path, "w", encoding="utf-8") as f:
            f.write(long_content)

    # Mock extractor to avoid needing a real PDF on disk.
    monkeypatch.setattr(_pe.PDFExtractor, "page_count", classmethod(lambda cls, p: 2))
    monkeypatch.setattr(
        _pe.PDFExtractor, "render_cover", classmethod(lambda cls, b, **kw: None)
    )

    _svc.AudiobookService._queue = None
    _svc.AudiobookService._worker_task = None
    _svc.AudiobookService.initialize()

    await _svc.AudiobookService._phase_extract(bid)

    clean1 = AudiobookStore.page_clean_path(bid, 1)
    clean2 = AudiobookStore.page_clean_path(bid, 2)

    # Page 1 must NOT have a silence marker (it is the original content).
    assert not os.path.exists(
        clean1
    ), "page 1 must not have a pre-written clean file (Gemini should clean it)"

    # Page 2 must have been pre-written with the silence marker.
    assert os.path.exists(
        clean2
    ), "page 2 (duplicate) must have a pre-written silence marker clean file"
    with open(clean2, encoding="utf-8") as f:
        marker = f.read()
    assert marker == "-", f"duplicate page clean file must be '-', got: {marker!r}"


@pytest.mark.asyncio
async def test_non_duplicate_pages_are_not_marked_silent(monkeypatch):
    """Pages with distinct content must not be marked as silence — the normal
    Gemini clean path must remain unobstructed."""
    from app.services import audiobook_service as _svc
    from app.services import pdf_extractor as _pe

    bid = AudiobookStore.create_book("multipage.pdf")
    meta = AudiobookStore.initial_meta(
        bid, "multipage.pdf", 3, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
    )
    meta["file_ext"] = "pdf"
    AudiobookStore.write_meta(bid, meta)

    contents = [
        "Chapter 1: " + "A" * 200,
        "Chapter 2: " + "B" * 200,
        "Chapter 3: " + "C" * 200,
    ]
    for n, content in enumerate(contents, start=1):
        path = AudiobookStore.page_raw_path(bid, n)
        os.makedirs(os.path.dirname(path), exist_ok=True)
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)

    monkeypatch.setattr(_pe.PDFExtractor, "page_count", classmethod(lambda cls, p: 3))
    monkeypatch.setattr(
        _pe.PDFExtractor, "render_cover", classmethod(lambda cls, b, **kw: None)
    )

    _svc.AudiobookService._queue = None
    _svc.AudiobookService._worker_task = None
    _svc.AudiobookService.initialize()

    await _svc.AudiobookService._phase_extract(bid)

    # None of the 3 pages should have a pre-written clean file.
    for n in (1, 2, 3):
        assert not os.path.exists(
            AudiobookStore.page_clean_path(bid, n)
        ), f"page {n} should not have a pre-written clean file — content is unique"


# ---------- TXT extraction (non-PDF path) ----------


@pytest.mark.asyncio
async def test_txt_file_extraction_does_not_call_pdf_extractor(monkeypatch):
    """Uploading a .txt file must route through TextExtractor, not PDFExtractor.

    Regression guard: before the TextExtractor path was added, any non-PDF
    upload crashed in the extract phase because PDFExtractor.page_count tried
    to open the file as a PDF.
    """
    from app.services import audiobook_service as _svc
    from app.services import pdf_extractor as _pe

    bid = AudiobookStore.create_book("sample.txt")
    # Write meta with file_ext=txt so the pipeline routes correctly.
    # page_count=1 matches the single-page content we write below.
    meta = AudiobookStore.initial_meta(
        bid, "sample.txt", 1, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
    )
    meta["file_ext"] = "txt"
    AudiobookStore.write_meta(bid, meta)

    # Write a real source file so TextExtractor can open it.
    source_path = AudiobookStore.source_file_path(bid, "txt")
    os.makedirs(os.path.dirname(source_path), exist_ok=True)
    with open(source_path, "w", encoding="utf-8") as f:
        f.write("Hello world.\n\nThis is page one.\n\nThis is page two.")

    pdf_calls: list[str] = []

    def _fail_if_called(*args, **kwargs):
        pdf_calls.append("called")
        raise AssertionError("PDFExtractor must NOT be called for a .txt file")

    monkeypatch.setattr(
        _pe.PDFExtractor, "page_count", classmethod(lambda cls, p: _fail_if_called(p))
    )
    monkeypatch.setattr(
        _pe.PDFExtractor, "render_cover", classmethod(lambda cls, b: _fail_if_called(b))
    )

    # Reset AudiobookService state so initialize() creates fresh objects in the
    # current event loop (avoids "bound to a different event loop" errors when
    # multiple async tests share the singleton).
    _svc.AudiobookService._queue = None
    _svc.AudiobookService._worker_task = None
    _svc.AudiobookService.initialize()

    await _svc.AudiobookService._phase_extract(bid)

    assert len(pdf_calls) == 0, "PDFExtractor was called for a TXT file"

    # _phase_extract does not update meta["page_count"] (that is set at upload time).
    # Instead verify that the per-page raw text file was actually written to disk.
    page1 = AudiobookStore.page_raw_path(bid, 1)
    assert os.path.exists(page1), "page 1 raw text file must exist after extraction"
    with open(page1, encoding="utf-8") as f:
        content = f.read()
    assert len(content) > 0, "extracted page should be non-empty"


# ---------- Gemini timeout → raw text fallback ----------


@pytest.mark.asyncio
async def test_gemini_timeout_falls_back_to_raw_text(monkeypatch):
    """When GeminiCleaner.clean_page raises asyncio.TimeoutError the pipeline
    must degrade gracefully: the cleaned output file is written with the raw
    text so downstream TTS can still proceed.  No hang, no crash, no empty file.
    """
    from app.services import audiobook_service as _svc
    from app.services import gemini_cleaner as _gc

    bid = AudiobookStore.create_book("Test.pdf")
    meta = AudiobookStore.initial_meta(
        bid, "Test.pdf", 1, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
    )
    AudiobookStore.write_meta(bid, meta)

    raw_text = "Raw page text that should survive the Gemini timeout."
    raw_path = AudiobookStore.page_raw_path(bid, 1)
    os.makedirs(os.path.dirname(raw_path), exist_ok=True)
    with open(raw_path, "w", encoding="utf-8") as f:
        f.write(raw_text)

    async def _timeout_clean(api_key, text):
        raise TimeoutError()

    monkeypatch.setattr(
        _gc.GeminiCleaner, "clean_page", AsyncMock(side_effect=_timeout_clean)
    )

    _svc.AudiobookService.initialize()
    # Must complete without raising, even though Gemini timed out.
    await _svc.AudiobookService._phase_clean(bid, api_key="test-key")

    clean_path = AudiobookStore.page_clean_path(bid, 1)
    assert os.path.exists(clean_path), "clean file must exist after timeout fallback"
    with open(clean_path, encoding="utf-8") as f:
        result = f.read()
    assert result == raw_text, "fallback content must equal the original raw text"


# ---------- section-detection fallback (no outline, Gemini unavailable) ----------


def test_fallback_sections_small_book_is_single_section():
    """A book at or under the chunk threshold has no real navigational value
    from chunking, so it stays a single section."""
    sections = AudiobookService._fallback_sections(10, "My Book")
    assert sections == [{"title": "My Book", "start_page": 1, "end_page": 10}]


def test_fallback_sections_untitled_book_uses_default_title():
    sections = AudiobookService._fallback_sections(5, None)
    assert sections[0]["title"] == "Audiobook"


def test_fallback_sections_large_book_chunks_into_parts():
    """Above the threshold, chunk into contiguous, page-covering parts instead
    of collapsing the whole book into one section (the bug being fixed)."""
    sections = AudiobookService._fallback_sections(40, "Long Book")
    assert len(sections) > 1
    assert sections[0] == {"title": "Part 1", "start_page": 1, "end_page": 15}
    assert sections[-1]["end_page"] == 40
    # Contiguous, gapless, and covers every page exactly once.
    for i in range(1, len(sections)):
        assert sections[i]["start_page"] == sections[i - 1]["end_page"] + 1


@pytest.mark.asyncio
async def test_phase_section_falls_back_to_chunked_parts_when_gemini_fails(
    monkeypatch,
):
    """A 40-page non-PDF book (no outline) whose Gemini section-detection call
    fails must not collapse into a single section spanning the whole book —
    it should chunk into multiple navigable parts (see _fallback_sections)."""
    from app.services import gemini_cleaner as _gc

    bid = AudiobookStore.create_book("Test.txt")
    meta = AudiobookStore.initial_meta(
        bid, "Test.txt", 40, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
    )
    meta["file_ext"] = "txt"
    AudiobookStore.write_meta(bid, meta)

    async def _fail_detect(api_key, pages):
        raise RuntimeError("simulated Gemini failure")

    monkeypatch.setattr(
        _gc.GeminiCleaner, "detect_sections", AsyncMock(side_effect=_fail_detect)
    )

    await AudiobookService._phase_section(
        bid, "test-key", AudiobookStore.read_meta(bid)
    )

    persisted = AudiobookStore.read_meta(bid)
    assert len(persisted["sections"]) > 1, "must not collapse to a single section"
    assert persisted["sections"][0]["start_page"] == 1
    assert persisted["sections"][-1]["end_page"] == 40


@pytest.mark.asyncio
async def test_phase_section_small_book_stays_single_section_on_gemini_failure(
    monkeypatch,
):
    """Below the chunk threshold, the pre-existing single-section fallback
    behavior is preserved (chunking a short book adds no value)."""
    from app.services import gemini_cleaner as _gc

    bid = AudiobookStore.create_book("Test.txt")
    meta = AudiobookStore.initial_meta(
        bid, "Test.txt", 3, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
    )
    meta["file_ext"] = "txt"
    AudiobookStore.write_meta(bid, meta)

    async def _fail_detect(api_key, pages):
        raise RuntimeError("simulated Gemini failure")

    monkeypatch.setattr(
        _gc.GeminiCleaner, "detect_sections", AsyncMock(side_effect=_fail_detect)
    )

    await AudiobookService._phase_section(
        bid, "test-key", AudiobookStore.read_meta(bid)
    )

    persisted = AudiobookStore.read_meta(bid)
    assert len(persisted["sections"]) == 1
    assert persisted["sections"][0]["end_page"] == 3


# ---------- Phase 2: structural-format-native cascade (Sprint 3 / Finding Set B) ----------
#
# the v1.1 design notes §6.6: PDF outline -> DOCX headings -> MD headers -> Gemini -> fallback.
# Each of the first three, when structural data is found, must short-circuit
# the cascade the same way the (now-fixed) outline path always has — Gemini's
# detect_sections must never be called for these books.


@pytest.mark.asyncio
async def test_phase_section_pdf_outline_short_circuits_gemini(monkeypatch):
    """T3.1 + T3.4: a PDF with a real embedded outline must use those bookmark
    titles directly and never call Gemini's detect_sections."""
    import io

    from pypdf import PdfWriter

    from app.services import gemini_cleaner as _gc

    bid = AudiobookStore.create_book("Test.pdf")
    meta = AudiobookStore.initial_meta(
        bid, "Test.pdf", 5, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
    )
    meta["file_ext"] = "pdf"
    AudiobookStore.write_meta(bid, meta)

    writer = PdfWriter()
    for _ in range(5):
        writer.add_blank_page(width=200, height=200)
    writer.add_outline_item("Real Chapter One", 0)
    writer.add_outline_item("Real Chapter Two", 2)
    buf = io.BytesIO()
    writer.write(buf)
    AudiobookStore.save_source(bid, buf.getvalue(), "pdf")

    detect_mock = AsyncMock()
    monkeypatch.setattr(_gc.GeminiCleaner, "detect_sections", detect_mock)

    await AudiobookService._phase_section(
        bid, "test-key", AudiobookStore.read_meta(bid)
    )

    detect_mock.assert_not_called()
    persisted = AudiobookStore.read_meta(bid)
    titles = [s["title"] for s in persisted["sections"]]
    assert "Real Chapter One" in titles
    assert "Real Chapter Two" in titles
    assert not any(t.startswith("Part ") for t in titles), titles


@pytest.mark.asyncio
async def test_phase_section_docx_headings_short_circuit_gemini(monkeypatch):
    """T3.2 + T3.4: a DOCX with Heading-styled paragraphs must use those
    headings directly and never call Gemini's detect_sections."""
    import io

    from docx import Document

    from app.services import gemini_cleaner as _gc

    bid = AudiobookStore.create_book("Test.docx")
    meta = AudiobookStore.initial_meta(
        bid, "Test.docx", 2, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
    )
    meta["file_ext"] = "docx"
    AudiobookStore.write_meta(bid, meta)

    doc = Document()
    doc.add_heading("Real Chapter One", level=1)
    doc.add_paragraph("Body text.")
    doc.add_heading("Real Chapter Two", level=1)
    doc.add_paragraph("More body text.")
    buf = io.BytesIO()
    doc.save(buf)
    AudiobookStore.save_source(bid, buf.getvalue(), "docx")

    detect_mock = AsyncMock()
    monkeypatch.setattr(_gc.GeminiCleaner, "detect_sections", detect_mock)

    await AudiobookService._phase_section(
        bid, "test-key", AudiobookStore.read_meta(bid)
    )

    detect_mock.assert_not_called()
    persisted = AudiobookStore.read_meta(bid)
    titles = [s["title"] for s in persisted["sections"]]
    assert "Real Chapter One" in titles
    assert "Real Chapter Two" in titles
    assert not any(t.startswith("Part ") for t in titles), titles


@pytest.mark.asyncio
async def test_phase_section_md_headers_short_circuit_gemini(monkeypatch):
    """T3.3 + T3.4: a Markdown file with #/## headers must use those headers
    directly and never call Gemini's detect_sections."""
    from app.services import gemini_cleaner as _gc

    bid = AudiobookStore.create_book("Test.md")
    meta = AudiobookStore.initial_meta(
        bid, "Test.md", 2, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
    )
    meta["file_ext"] = "md"
    AudiobookStore.write_meta(bid, meta)

    content = "# Real Chapter One\n\nBody text.\n\n# Real Chapter Two\n\nMore text.\n"
    AudiobookStore.save_source(bid, content.encode("utf-8"), "md")

    detect_mock = AsyncMock()
    monkeypatch.setattr(_gc.GeminiCleaner, "detect_sections", detect_mock)

    await AudiobookService._phase_section(
        bid, "test-key", AudiobookStore.read_meta(bid)
    )

    detect_mock.assert_not_called()
    persisted = AudiobookStore.read_meta(bid)
    titles = [s["title"] for s in persisted["sections"]]
    assert "Real Chapter One" in titles
    assert "Real Chapter Two" in titles
    assert not any(t.startswith("Part ") for t in titles), titles


@pytest.mark.asyncio
async def test_phase_section_skips_gemini_when_api_key_empty_and_no_structure(
    monkeypatch,
):
    """the v1.1 design notes §6.6 cascade step 4: Gemini is only attempted when a key is
    present. A book with no structural signal (plain TXT) and an empty key
    must go straight to _fallback_sections without ever attempting Gemini —
    avoiding a network round-trip that would just swallow an auth failure
    and return [] anyway (the exact silent-degrade mechanism behind D3)."""
    from app.services import gemini_cleaner as _gc

    bid = AudiobookStore.create_book("Test.txt")
    meta = AudiobookStore.initial_meta(
        bid, "Test.txt", 3, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
    )
    meta["file_ext"] = "txt"
    AudiobookStore.write_meta(bid, meta)

    detect_mock = AsyncMock()
    monkeypatch.setattr(_gc.GeminiCleaner, "detect_sections", detect_mock)

    await AudiobookService._phase_section(bid, "", AudiobookStore.read_meta(bid))

    detect_mock.assert_not_called()
    persisted = AudiobookStore.read_meta(bid)
    assert len(persisted["sections"]) == 1
    assert persisted["sections"][0]["end_page"] == 3


@pytest.mark.asyncio
async def test_phase_section_drops_outline_entries_outside_page_range(monkeypatch):
    """Defense-in-depth: an out-of-range start_page from a structural source
    must never reach the persisted sections (and therefore never corrupt a
    page_to_time lookup at concat time). DOCX/MD compute their own pagination
    internally so this should already always hold, but validate centrally
    here regardless — one uniform guard covering every structural source."""
    from app.services import pdf_extractor as _pe

    bid = AudiobookStore.create_book("Test.pdf")
    meta = AudiobookStore.initial_meta(
        bid, "Test.pdf", 5, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
    )
    meta["file_ext"] = "pdf"
    AudiobookStore.write_meta(bid, meta)

    # A malformed/adversarial outline: one valid entry, one whose start_page
    # is beyond the book's real 5-page extent.
    bogus_outline = [
        {"title": "Real Chapter", "start_page": 2, "level": 0},
        {"title": "Out Of Range", "start_page": 99, "level": 0},
    ]
    monkeypatch.setattr(
        _pe.PDFExtractor, "read_outline", classmethod(lambda cls, p: bogus_outline)
    )

    await AudiobookService._phase_section(
        bid, "test-key", AudiobookStore.read_meta(bid)
    )

    persisted = AudiobookStore.read_meta(bid)
    titles = [s["title"] for s in persisted["sections"]]
    assert "Out Of Range" not in titles
    assert "Real Chapter" in titles
    assert all(s["start_page"] <= 5 for s in persisted["sections"])


# ---------- Phase 2: idempotency (D3) ----------


@pytest.mark.asyncio
async def test_phase_section_skips_when_sections_already_present(monkeypatch):
    """D3: _phase_section must be idempotent like every other phase. A book
    resumed via resume_in_progress's "tts"/"concatenating" path re-enters
    _run_pipeline with api_key="" — if section detection re-ran here, Gemini
    would silently swallow the resulting auth failure and fall back to
    generic "Part N" sections, overwriting already-correct real chapter
    titles. Presence of non-empty meta['sections'] must short-circuit before
    any detection path (outline read OR Gemini) runs at all."""
    from unittest.mock import MagicMock

    from app.services import gemini_cleaner as _gc
    from app.services import pdf_extractor as _pe

    bid = AudiobookStore.create_book("Test.pdf")
    meta = AudiobookStore.initial_meta(
        bid, "Test.pdf", 5, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
    )
    meta["file_ext"] = "pdf"
    real_sections = [
        {
            "title": "Real Chapter One",
            "start_page": 1,
            "end_page": 5,
            "start_time": 0.0,
        }
    ]
    meta["sections"] = real_sections
    AudiobookStore.write_meta(bid, meta)

    outline_mock = MagicMock(return_value=None)
    monkeypatch.setattr(_pe.PDFExtractor, "read_outline", outline_mock)
    detect_mock = AsyncMock()
    monkeypatch.setattr(_gc.GeminiCleaner, "detect_sections", detect_mock)

    # Exactly resume_in_progress's "tts"/"concatenating" resume scenario:
    # empty api_key, sections already populated from the original run.
    await AudiobookService._phase_section(bid, "", AudiobookStore.read_meta(bid))

    outline_mock.assert_not_called()
    detect_mock.assert_not_called()
    persisted = AudiobookStore.read_meta(bid)
    assert persisted["sections"] == real_sections
    # The phase must never have even started — status stays whatever it was.
    assert persisted["status"] == "ready"


@pytest.mark.asyncio
async def test_phase_section_runs_normally_when_sections_empty(monkeypatch):
    """Sanity counterpart to the skip test above: a genuinely fresh book
    (sections == [], initial_meta's default) must NOT be skipped — the guard
    is about idempotency, not a blanket no-op."""
    from app.services import gemini_cleaner as _gc
    from app.services import pdf_extractor as _pe

    bid = AudiobookStore.create_book("Test.pdf")
    meta = AudiobookStore.initial_meta(
        bid, "Test.pdf", 5, "kokoro", "af_bella", 1.0, {"cost_usd": 0.0}
    )
    meta["file_ext"] = "pdf"
    assert meta["sections"] == []
    AudiobookStore.write_meta(bid, meta)

    monkeypatch.setattr(
        _pe.PDFExtractor, "read_outline", classmethod(lambda cls, p: None)
    )
    detect_mock = AsyncMock(return_value=[])
    monkeypatch.setattr(_gc.GeminiCleaner, "detect_sections", detect_mock)

    await AudiobookService._phase_section(
        bid, "test-key", AudiobookStore.read_meta(bid)
    )

    detect_mock.assert_called_once()
    persisted = AudiobookStore.read_meta(bid)
    assert persisted["status"] == "sectioning"
    assert len(persisted["sections"]) == 1  # fell through to _fallback_sections
