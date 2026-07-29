"""TextExtractor — TXT / DOCX / Markdown → audiobook page pipeline.

Mirrors the PDFExtractor interface so audiobook_service.py can route to
either class based on meta["file_ext"] without branching everywhere.

All file-system writes are atomic (tmp+rename). All methods are sync;
callers wrap in run_in_executor when called from async code.
"""

import io
import os
import re
import uuid

from PIL import Image, ImageDraw

from app.core.logging import get_logger
from app.services.audiobook_store import AudiobookStore

log = get_logger(__name__)

# Target words per synthetic page. 400 words ≈ 2-3 minutes of audio.
_WORDS_PER_PAGE = 400


class TextExtractor:
    # ---------- text reading ----------

    @classmethod
    def read_text(cls, source_path: str) -> str:
        """Return the full plain-text content of a TXT, MD, or DOCX file."""
        ext = os.path.splitext(source_path)[1].lower()
        if ext == ".docx":
            from docx import Document  # python-docx

            doc = Document(source_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)
        else:
            with open(source_path, encoding="utf-8", errors="replace") as f:
                return f.read()

    @classmethod
    def split_pages(cls, text: str) -> list[str]:
        """Split text into ~_WORDS_PER_PAGE-word pages at paragraph boundaries.

        A paragraph that alone exceeds the page budget — including the whole
        document when it has no blank-line breaks at all — is chunked by raw
        word count instead of becoming one oversized page. The undersized
        remainder feeds into normal accumulation with whatever paragraph
        follows, so it doesn't strand a near-empty trailing page.
        """
        pages, _ = cls._paginate(cls._split_into_paragraphs(text))
        return pages

    @staticmethod
    def _split_into_paragraphs(text: str) -> list[str]:
        """Normalise line endings and split on blank lines into paragraphs."""
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        raw_paras = re.split(r"\n{2,}", text)
        return [p.strip() for p in raw_paras if p.strip()]

    @classmethod
    def _paginate(cls, paragraphs: list[str]) -> tuple[list[str], list[int]]:
        """Accumulate paragraphs into ~_WORDS_PER_PAGE-word pages.

        Returns `(pages, paragraph_start_page)`: `pages` is identical to what
        `split_pages` has always returned; `paragraph_start_page[i]` is the
        1-indexed synthetic page on which `paragraphs[i]` BEGINS. Structural
        heading detection (DOCX heading styles, Markdown `#`/`##`) reuses this
        so a heading's reported `start_page` always matches the same page
        boundaries TTS/extraction actually use — no separate, driftable
        pagination logic.
        """
        pages: list[str] = []
        paragraph_start_page: list[int] = []
        current: list[str] = []
        current_words = 0

        for para in paragraphs:
            words = para.split()
            wc = len(words)
            if current_words + wc > _WORDS_PER_PAGE and current:
                pages.append("\n\n".join(current))
                current = []
                current_words = 0
            # The page this paragraph starts on: the next page to be opened
            # (pages already flushed above, plus the one `current` becomes).
            paragraph_start_page.append(len(pages) + 1)
            if wc > _WORDS_PER_PAGE:
                i = 0
                while wc - i > _WORDS_PER_PAGE:
                    pages.append(" ".join(words[i : i + _WORDS_PER_PAGE]))
                    i += _WORDS_PER_PAGE
                tail = words[i:]
                current = [" ".join(tail)] if tail else []
                current_words = len(tail)
            else:
                current.append(para)
                current_words += wc

        if current:
            pages.append("\n\n".join(current))

        return (pages if pages else [""]), paragraph_start_page

    # ---------- PDFExtractor-compatible interface ----------

    @classmethod
    def page_count(cls, source_path: str) -> int:
        text = cls.read_text(source_path)
        return len(cls.split_pages(text))

    @classmethod
    def is_image_only(cls, source_path: str) -> bool:
        return False

    @classmethod
    def sample_word_count(cls, source_path: str) -> int:
        text = cls.read_text(source_path)
        pages = cls.split_pages(text)
        n = len(pages)
        if n == 0:
            return 0
        indices = sorted({0, n // 2, n - 1})
        samples = [len(pages[i].split()) for i in indices]
        return sum(samples) // len(samples)

    @classmethod
    def sample_char_count(cls, source_path: str) -> int:
        text = cls.read_text(source_path)
        pages = cls.split_pages(text)
        n = len(pages)
        if n == 0:
            return 0
        indices = sorted({0, n // 2, n - 1})
        samples = [len(pages[i]) for i in indices]
        return sum(samples) // len(samples)

    @classmethod
    def extract_one(cls, book_id: str, page_num: int) -> None:
        """Write page_num (1-indexed) to pages/{n:03d}.txt.

        On first call for a book, splits and writes ALL pages at once so
        subsequent calls for pages 2..N find their files and skip I/O.
        """
        out = AudiobookStore.page_raw_path(book_id, page_num)
        if os.path.exists(out):
            return

        meta = AudiobookStore.read_meta(book_id) or {}
        file_ext = meta.get("file_ext", "txt")
        source_path = AudiobookStore.source_file_path(book_id, file_ext)

        text = cls.read_text(source_path)
        pages = cls.split_pages(text)

        for i, page_text in enumerate(pages, start=1):
            p = AudiobookStore.page_raw_path(book_id, i)
            if not os.path.exists(p):
                cls._atomic_write(p, page_text)

    @classmethod
    def read_outline(cls, source_path: str) -> list[dict] | None:
        """Return structural sections from DOCX/Markdown when they exist.

        The page positions reuse the exact synthetic pagination that feeds TTS,
        preventing section navigation from drifting away from audio playback.
        Plain text has no structural heading signal and returns ``None``.
        """
        extension = os.path.splitext(source_path)[1].lower()
        if extension == ".docx":
            return cls._docx_headings(source_path)
        if extension == ".md":
            return cls._markdown_headings(source_path)
        return None

    @classmethod
    def _docx_headings(cls, source_path: str) -> list[dict] | None:
        """Map DOCX Heading N styles to the page where narration reaches them."""
        try:
            from docx import Document

            document = Document(source_path)
            paragraphs = [
                paragraph for paragraph in document.paragraphs if paragraph.text.strip()
            ]
            if not paragraphs:
                return None
            pipeline_paragraphs = cls._split_into_paragraphs(cls.read_text(source_path))
            if not pipeline_paragraphs:
                return None
            _, start_pages = cls._paginate(pipeline_paragraphs)

            headings: list[dict] = []
            pipeline_index = 0
            for paragraph in paragraphs:
                parts = cls._split_into_paragraphs(paragraph.text) or [
                    paragraph.text.strip()
                ]
                style_name = (
                    paragraph.style.name if paragraph.style is not None else ""
                ) or ""
                if style_name.startswith("Heading"):
                    suffix = style_name[len("Heading") :].strip()
                    level = int(suffix) if suffix.isdigit() else 1
                    headings.append(
                        {
                            "title": parts[0],
                            "start_page": start_pages[
                                min(pipeline_index, len(start_pages) - 1)
                            ],
                            "level": level,
                        }
                    )
                pipeline_index += len(parts)
            return headings or None
        except Exception as error:
            log.warning(
                "docx.heading_read_failed",
                extra={"path": source_path, "error": str(error)},
            )
            return None

    @classmethod
    def _markdown_headings(cls, source_path: str) -> list[dict] | None:
        """Map level-one and level-two ATX Markdown headings to TTS pages."""
        try:
            with open(source_path, encoding="utf-8", errors="replace") as file:
                paragraphs = cls._split_into_paragraphs(file.read())
            if not paragraphs:
                return None
            _, start_pages = cls._paginate(paragraphs)
            headings: list[dict] = []
            for paragraph, start_page in zip(paragraphs, start_pages):
                parsed = cls._parse_md_header(paragraph.split("\n", 1)[0])
                if parsed is not None:
                    level, title = parsed
                    headings.append(
                        {"title": title, "start_page": start_page, "level": level}
                    )
            return headings or None
        except OSError as error:
            log.warning(
                "markdown.heading_read_failed",
                extra={"path": source_path, "error": str(error)},
            )
            return None

    @staticmethod
    def _parse_md_header(line: str) -> tuple[int, str] | None:
        """Parse a level-one/two CommonMark ATX heading without harming ``F#``."""
        stripped = line.strip()
        hashes = len(stripped) - len(stripped.lstrip("#"))
        if (
            hashes < 1
            or hashes > 2
            or hashes >= len(stripped)
            or stripped[hashes] != " "
        ):
            return None
        title = stripped[hashes:].strip()
        without_closer = title.rstrip("#")
        if without_closer != title and (
            not without_closer or without_closer[-1] == " "
        ):
            title = without_closer.rstrip()
        return (hashes, title) if title else None

    # ---------- cover ----------

    @classmethod
    def render_cover(cls, book_id: str) -> None:
        """Generate a minimal placeholder cover JPEG. Skip if already exists."""
        out = AudiobookStore.cover_path(book_id)
        if os.path.exists(out):
            return

        meta = AudiobookStore.read_meta(book_id) or {}
        file_ext = (meta.get("file_ext") or "txt").upper()

        img = Image.new("RGB", (600, 840), color=(18, 26, 38))
        draw = ImageDraw.Draw(img)

        # Cyan accent bar
        draw.rectangle([0, 0, 600, 10], fill=(0, 210, 230))

        # File-type badge (bottom-right)
        badge_text = f".{file_ext}"
        badge_x, badge_y = 470, 760
        draw.rectangle(
            [badge_x - 10, badge_y - 6, badge_x + 100, badge_y + 26], fill=(0, 180, 200)
        )
        draw.text((badge_x, badge_y), badge_text, fill=(255, 255, 255))

        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=85, optimize=True)
        cls._atomic_write_bytes(out, buf.getvalue())

    # ---------- atomic helpers ----------

    @staticmethod
    def _atomic_write(path: str, text: str) -> None:
        os.makedirs(os.path.dirname(path), exist_ok=True)
        tmp = path + ".tmp"
        with open(tmp, "w", encoding="utf-8") as f:
            f.write(text)
        os.replace(tmp, path)

    @staticmethod
    def _atomic_write_bytes(path: str, data: bytes) -> None:
        """D7: the tmp suffix is unique per call (not a fixed `path + ".tmp"`).

        Two independent call sites can render the same book's cover.jpg —
        api/audiobook.py's upload-time background task and
        _phase_extract's own render — both idempotency-gated by an
        `if os.path.exists(out): return` check with a TOCTOU window. A fixed
        shared tmp path let concurrent writers interleave writes into the
        same inode (or race each other's os.replace) and corrupt cover.jpg.
        With a unique tmp path per call, concurrent writers never touch the
        same file before either os.replace()s — which is itself atomic — so
        the loser's replace simply overwrites the winner's with an equally
        complete, valid image; never a half-written mix of both.
        """
        os.makedirs(os.path.dirname(path), exist_ok=True)
        tmp = f"{path}.{uuid.uuid4().hex}.tmp"
        with open(tmp, "wb") as f:
            f.write(data)
        os.replace(tmp, path)
